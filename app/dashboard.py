import os
import io
import time
import asyncio
import streamlit as st
import pandas as pd
import plotly.express as px
import seaborn as sns
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches
from config import SAVED_RUNS_DIR, SAVED_RUNS_DIR_PP, PP
from modules.scrapers.scraper import scrape_marktplaats_playwright
from modules.scrapers.scraper_2dehands import scrape_2dehands_playwright
from modules.scrapers.scrape_vinted import scrape_vinted_playwright
from modules.scrapers.scrape_ebay import scrape_ebay_playwright
from modules.scrapers.scrape_catawiki import scrape_catawiki_playwright
from modules.scrapers.scraper_markt_de import scrape_markt_playwright
from modules.query_asssistant import show_query_assistant
from modules import image_search
import os
import sys
import json
import logging
import threading
import numpy as np
import datetime
import traceback
from pathlib import Path
import base64
import uuid

# Global API keys
openai_api_key = os.getenv("OPENAI_API_KEY")
brevo_api_key = os.getenv("BREVO_API_KEY")
groq_api_key = os.getenv("GROQ_API_KEY")
deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")

# Setup directory paths
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)  # Go up one level from app directory
# Define raw data directory
RAW_DATA_DIR = os.path.join(project_root, "data", "raw_data")
os.makedirs(RAW_DATA_DIR, exist_ok=True)

# Define the directory to store saved CSV files
SAVED_RUNS_DIR = os.path.join(project_root, "data", "raw_data")
os.makedirs(SAVED_RUNS_DIR, exist_ok=True)

SAVED_RUNS_DIR_PP = os.path.join(project_root, "data", "pp")
os.makedirs(SAVED_RUNS_DIR_PP, exist_ok=True)

PP = os.path.join(project_root, "data", "pp_data") 
os.makedirs(PP, exist_ok=True)

# File paths
USER_AGENTS_PATH = os.path.join(current_dir, "data", "user-agents.txt")
CAT_FILE_PATH = os.path.join(current_dir, "data", "cat.txt")
CSV_EXPORT_PATH = os.path.join(project_root, "data", "class_data", "{search_query_encoded}.csv")

# Add project root to path for module imports
if project_root not in sys.path:
    sys.path.append(project_root) 

# Database configuration
database_url = os.getenv("DATABASE_URL")
if not database_url:
    raise ValueError("DATABASE_URL environment variable must be set")

# Supabase configuration
supabase_url = os.getenv("SUPABASE_URL")
supabase_key = os.getenv("SUPABASE_KEY")
if not supabase_url or not supabase_key:
    raise ValueError("SUPABASE_URL and SUPABASE_KEY environment variables must be set")

# Updating imports to use the new redundant storage system
from modules.db import (
    save_raw_data_redundant, 
    save_enriched_data_redundant,
    save_preprocessed_data_redundant,
    start_file_watcher,
    get_latest_listings
)
from modules.db.ps.database import create_tables_if_not_exists as create_postgres_tables

from app.classification import enrich_dataframe
from analysis import (
    run_agent_analysis,
    sentiment_analysis_agent,
    detailed_comparative_analysis,
    price_trend_analysis,
    market_insights_analysis,
    best_deal_recommendation,
    market_demand_forecasting
)
from email_utils import send_email_insights
from utils import (
 clean_price,
  get_category_options,
    get_raw_data_files,
    get_class_data_files,
    is_valid_email,
    sanitize_filename_component,
    preprocess_enriched_df,
)



st.set_page_config(
    page_title="Discovr",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for styling the app
st.markdown("""
<style>
    div[data-testid="stAppViewContainer"]{
        margin-right:2rem;
        margin-top:4rem;
        margin-left:2rem;
        margin-bottom:4rem;
        border-radius:2rem;
        border: 2px solid #e0e4e8;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
    }
    div[data-testid="stSidebarContent"] {
        background-color: #ffffff;
        border-right: 1px solid #e0e4e8;
    }
    div[data-testid="stSidebarCollapsedControl"]{
        margin-top:4rem;
        margin-left:2rem;
    }
    header[data-testid="stHeader"]{
        background-color:transparent;
        margin-top:2rem;
        margin-right:2rem;
    }
    .main .block-container {
        padding: 2rem;
        border-radius: 1rem;
        background-color: white;
        border: 1px solid #e0e4e8;
    }
</style>
""", unsafe_allow_html=True)



def main():
   
    st.title("🔍 Discovr")
    
    # Initialize session state variables if they don't exist
    if "page" not in st.session_state:
        st.session_state.page = "getting_data"
    
    st.markdown("""
    **Unparalleled Insights:** Data-Driven Analysis Delivered Directly to Your Inbox.

    This application provides a comprehensive approach to market analysis. By gathering data from online marketplaces through web scraping, 
    we process and analyze this information using advanced techniques.

    Our algorithms uncover valuable insights into market trends, pricing, and consumer sentiment. 
    These insights are compiled and presented in an easy-to-understand format, allowing you to make informed decisions.

    **Additionally**, we deliver these insights directly to your inbox, ensuring you never miss important market developments. 
    Stay informed, stay competitive, and let our data-driven analysis be your guide to success.
    """)

    # Create sidebar for navigation
    st.sidebar.title("Navigation")
    
    # Define the available pages with icons and keys
    navigation_options = {
        "Extract Data": "getting_data",
        "Transform Data": "enrich_data",
        "Load data": "database_stats",
        "Dashboard": "dashboard",
        "Analytics": "analytics",
        "Comprehensive Analysis": "comprehensive",
        "Reports": "reports"
    }
    
    # Add custom CSS to style the navigation buttons
    st.markdown("""
    <style>
        div.stButton > button {
            width: 100%;
            text-align: left;
            padding: 0.5rem 0.5rem;
            border: none;
            border-radius: 0.5rem;
            background-color: #f0f2f6;
            margin-bottom: 0.5rem;
            font-weight: 700;
        }
        div.stButton > button:hover {
            background-color: #e0e4e8;
            color: black;
        }
        .nav-button {
            display: flex;
            align-items: center;
        }
        
        
    </style>
    """, unsafe_allow_html=True)
    
    # Create buttons for each navigation option
    for page_name, page_key in navigation_options.items():
        # Check if this is the active page
        is_active = st.session_state.get("page") == page_key
        
        # Create a button for navigation
        with st.sidebar:
            if st.button(page_name, key=f"nav_{page_key}"):
                st.session_state["page"] = page_key
    # Display a divider in the sidebar
    st.sidebar.markdown("---")
    
    # Add additional sidebar information or controls
    st.sidebar.info(
        "This application allows you to collect, enrich, and analyze data from various online marketplaces."
    )
    
    # Page 1: Getting Data
    if st.session_state.page == "getting_data":
        st.header("Extract Data")

        # Website Selection
        website_options = {
            "Marktplaats": "scrape_marktplaats_playwright",
            "2dehands": "scrape_2dehands_playwright", 
            "Vinted": "scrape_vinted_playwright",
            "eBay": "scrape_ebay_playwright",
            "Catawiki": "scrape_catawiki_playwright",
            "Markt.de": "scrape_markt_playwright"
        }
        
        selected_websites = st.multiselect(
            "Select Websites to Scrape",
            options=list(website_options.keys()),
            default=["Marktplaats"],
            key="website_selector"
        )

        # Search Assistant
        with st.expander("Use Search Assistant"):
            chosen_query = show_query_assistant()
            if chosen_query:
                st.session_state["assistant_query"] = chosen_query
                st.success(f"Chosen search query from assistant: {chosen_query}")

        with st.expander("Use Image Search"):
            image_search.find_similar_products()

        # Main scraping form
        with st.form(key='scrape_form'):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                # If we want to use the assistant query
                default_query = st.session_state.get("assistant_query", "")
                search_query = st.text_input("✅ Enter your search term:", value=default_query)

                # Only show category for websites that use it
                if any(site in ["Marktplaats", "2dehands"] for site in selected_websites):
                    category = st.selectbox("📂 Category", get_category_options())
                else:
                    category = ""  # Default empty for sites that don't use categories
                    
                max_pages = st.number_input("📄 Number of pages (1-200)", min_value=1, max_value=200, value=2)
            
            with col2:
                st.empty()  # Placeholder for alignment

            submit_scrape = st.form_submit_button(label="Search")
            estimated_time = (max_pages * 7 * len(selected_websites)) / 60
            st.info(f"⏱️ Estimated wait time: {estimated_time:.1f} minutes")

        if submit_scrape:
            if not search_query:
                st.error("Please enter a search query")
                return

            with st.spinner(f"🕵️‍♂️ Scraping selected websites..."):
                try:
                    # Import asyncio inside the function scope
                    import asyncio
                    all_results = []
                    
                    # Process each website one by one
                    for website in selected_websites:
                        st.info(f"Scraping {website}...")
                        
                        try:
                            # Create async function to handle scraping
                            async def do_scrape():
                                st.write(f"Starting async scraping for {website}...")
                                try:
                                    if website == "Marktplaats":
                                        result = await scrape_marktplaats_playwright(search_query, category, max_pages)
                                        st.write(f"Marktplaats scraper returned: {type(result)}")
                                        
                                        # Handle case where Marktplaats scraper returns a tuple
                                        if isinstance(result, tuple):
                                            st.write(f"Tuple length: {len(result)}")
                                            # Assuming the first element is the DataFrame we want
                                            if len(result) > 0 and isinstance(result[0], pd.DataFrame):
                                                df = result[0]
                                                st.write(f"Extracted DataFrame with {len(df)} rows")
                                                
                                                # Save related searches if available
                                                if len(result) > 1 and isinstance(result[1], pd.DataFrame):
                                                    related_df = result[1]
                                                    st.write(f"Also extracted related searches DataFrame with {len(related_df)} rows")
                                                    
                                                    # Store related searches in PostgreSQL
                                                    try:
                                                        from modules.db.ps.database import store_related_searches
                                                        
                                                        # Store in database
                                                        store_result = store_related_searches(
                                                            related_df=related_df,
                                                            original_search_query=search_query,
                                                            website=website,
                                                            category=category,
                                                            source_file=None
                                                        )
                                                        
                                                        if store_result["success"]:
                                                            st.success(f"✅ {store_result['inserted_count']} related searches stored in PostgreSQL database")
                                                        else:
                                                            st.warning(f"⚠️ Could not store related searches in database: {store_result.get('error', 'Unknown error')}")
                                                            
                                                            # Fall back to CSV if database storage fails
                                                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                                            search_query_encoded = sanitize_filename_component(search_query)
                                                            related_csv_filename = f"Related_{website}_{search_query_encoded}_{timestamp}.csv"
                                                            related_csv_path = os.path.join(SAVED_RUNS_DIR, related_csv_filename)
                                                            related_df.to_csv(related_csv_path, index=False)
                                                            st.info(f"📁 Related searches saved to CSV as fallback at: {related_csv_path}")
                                                    except Exception as save_e:
                                                        st.error(f"❌ Error storing related searches: {save_e}")
                                                
                                                return df
                                            else:
                                                st.error("Unable to extract DataFrame from tuple")
                                                return pd.DataFrame()
                                        elif isinstance(result, pd.DataFrame):
                                            st.write(f"Marktplaats returned DataFrame with {len(result)} rows")
                                            return result
                                        else:
                                            st.error(f"Unexpected return type from Marktplaats: {type(result)}")
                                            return pd.DataFrame()
                                    elif website == "2dehands":
                                        result = await scrape_2dehands_playwright(search_query, category, max_pages)
                                        st.write(f"2dehands scraper returned data: {not result.empty if isinstance(result, pd.DataFrame) else 'Not a DataFrame'}")
                                        return result
                                    elif website == "Vinted":
                                        result = await scrape_vinted_playwright(search_query=search_query, max_pages=max_pages)
                                        st.write(f"Vinted scraper returned data: {not result.empty if isinstance(result, pd.DataFrame) else 'Not a DataFrame'}")
                                        return result
                                    elif website == "eBay":
                                        result = await scrape_ebay_playwright(search_query=search_query, max_pages=max_pages)
                                        st.write(f"eBay scraper returned data: {not result.empty if isinstance(result, pd.DataFrame) else 'Not a DataFrame'}")
                                        return result
                                    elif website == "Catawiki":
                                        result = await scrape_catawiki_playwright(search_query=search_query, max_pages=max_pages)
                                        st.write(f"Catawiki scraper returned data: {not result.empty if isinstance(result, pd.DataFrame) else 'Not a DataFrame'}")
                                        return result
                                    elif website == "Markt.de":
                                        result = await scrape_markt_playwright(search_query=search_query, max_pages=max_pages)
                                        st.write(f"Markt.de scraper returned data: {not result.empty if isinstance(result, pd.DataFrame) else 'Not a DataFrame'}")
                                        return result
                                except Exception as scrape_e:
                                    st.error(f"Error in {website} scraper: {scrape_e}")
                                    return pd.DataFrame()  # Return empty DataFrame on error
                            
                            # Run the async function
                            st.write(f"Running async function for {website}...")
                            result = asyncio.run(do_scrape())
                            st.write(f"Async function completed for {website}")
                            
                            # Convert the result to a DataFrame if it isn't already
                            if isinstance(result, pd.DataFrame):
                                if result.empty:
                                    st.warning(f"Received empty DataFrame from {website}")
                                else:
                                    st.success(f"Received {len(result)} rows from {website}")
                                    result['source_file'] = website
                                    result['website'] = website
                                    all_results.append(result)
                                    st.success(f"Added {website} results to all_results (now contains {len(all_results)} DataFrames)")
                            else:
                                st.warning(f"Result from {website} is not a DataFrame: {type(result)}")
                            
                            st.success(f"Completed processing {website}")
                            
                        except Exception as e:
                            st.error(f"Error processing {website}: {e}")
                            continue

                    # Combine all results if there are multiple websites
                    st.write(f"Processing {len(all_results)} result DataFrames")
                    if len(all_results) > 1:
                        try:
                            df = pd.concat(all_results, ignore_index=True)
                            st.success(f"Combined {len(all_results)} DataFrames with a total of {len(df)} rows")
                        except Exception as e:
                            st.error(f"Error combining results: {e}")
                            # Let's try to understand what's in all_results
                            for i, res in enumerate(all_results):
                                st.write(f"DataFrame {i+1} type: {type(res)}, empty: {res.empty if isinstance(res, pd.DataFrame) else 'Not a DataFrame'}")
                            return
                    elif len(all_results) == 1:
                        df = all_results[0]
                        st.success(f"Using single DataFrame with {len(df)} rows")
                    else:
                        st.error("No results found. Please check the debugging information above to see which scrapers failed.")
                        return

                    if df.empty:
                        st.error("Combined DataFrame is empty. Please try a different search query or website.")
                        return

                    # Save raw data and store in databases
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    search_query_encoded = sanitize_filename_component(search_query)
                    
                    # Create filenames for each selected website
                    for website in selected_websites:
                        # Filter data for current website
                        website_df = df[df['source_file'] == website].copy()
                        
                        if not website_df.empty:
                            csv_filename = f"Raw_{website}_{search_query_encoded}_{timestamp}.csv"
                            csv_path = os.path.join(SAVED_RUNS_DIR, csv_filename)
                            
                            # Add metadata to the DataFrame before saving
                            website_df['search_query'] = search_query
                            website_df['category'] = category
                            website_df['website'] = website
                            
                            # Use redundant storage to save to both PostgreSQL and Supabase
                            try:
                                st.info(f"Saving data to databases for {website}...")
                                result = save_raw_data_redundant(website_df, source_file=csv_filename)
                                
                                if result["overall_success"]:
                                    if result["postgres"]["success"] and result["supabase"]["success"]:
                                        st.success(f"✅ Data successfully stored in both PostgreSQL and Supabase for {website}")
                                    elif result["postgres"]["success"]:
                                        st.warning(f"✅ Data stored in PostgreSQL but not Supabase for {website}")
                                    elif result["supabase"]["success"]:
                                        st.warning(f"✅ Data stored in Supabase but not PostgreSQL for {website}")
                                else:
                                    st.error(f"❌ Failed to store data in any database for {website}")
                                    
                                    # Save to CSV as final fallback
                                    try:
                                        website_df.to_csv(csv_path, index=False)
                                        st.info(f"📁 Data saved to CSV for {website} at: {csv_path}")
                                    except Exception as csv_e:
                                        st.error(f"❌ Error saving data to CSV for {website}: {csv_e}")
                            except Exception as e:
                                st.error(f"❌ Error during database operations: {e}")
                                
                                # Save to CSV as fallback
                                try:
                                    website_df.to_csv(csv_path, index=False)
                                    st.info(f"📁 Data saved to CSV for {website} at: {csv_path}")
                                except Exception as csv_e:
                                    st.error(f"❌ Error saving data to CSV for {website}: {csv_e}")

                    # Store in session state
                    st.session_state["raw_df"] = df
                    st.session_state["preprocessed_df"] = df

                    # Display results
                    st.subheader("Search Results")
                    st.dataframe(df)

                    # Start the file watcher
                    start_file_watcher()

                except Exception as e:
                    st.error(f"❌ Error during scraping: {e}")
                    return
    elif st.session_state.page == "database_stats":
        st.header("📊 Database Statistics")
        
        # Create tabs for different statistics views
        stats_tabs = st.tabs(["Storage Stats", "Data Distribution", "Performance Metrics", "Related Searches", "Search Listings"])
        
        with stats_tabs[0]:
            st.subheader("Database Storage Statistics")
            
            if st.button("Get Storage Stats", key="get_storage_stats"):
                try:
                    # PostgreSQL stats
                    from modules.db.ps.database import get_db_connection
                    conn = get_db_connection()
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("### PostgreSQL Stats")
                        
                        with conn.cursor() as cur:
                            # Raw listings count
                            cur.execute("SELECT COUNT(*) FROM raw_listings")
                            raw_count = cur.fetchone()[0]
                            
                            # Enriched listings count
                            cur.execute("SELECT COUNT(*) FROM enriched_listings")
                            enriched_count = cur.fetchone()[0]
                            
                            # Preprocessed listings count
                            cur.execute("SELECT COUNT(*) FROM preprocessed_listings")
                            preprocessed_count = cur.fetchone()[0]
                            
                            # Table sizes
                            cur.execute("""
                                SELECT 
                                    table_name, 
                                    pg_size_pretty(pg_total_relation_size(quote_ident(table_name)))
                                FROM 
                                    information_schema.tables
                                WHERE 
                                    table_schema = 'public'
                                    AND table_name IN ('raw_listings', 'enriched_listings', 'preprocessed_listings')
                            """)
                            table_sizes = cur.fetchall()
                        
                        st.markdown(f"**Raw Listings:** {raw_count:,}")
                        st.markdown(f"**Enriched Listings:** {enriched_count:,}")
                        st.markdown(f"**Preprocessed Listings:** {preprocessed_count:,}")
                        
                        st.markdown("**Table Sizes:**")
                        for table, size in table_sizes:
                            st.markdown(f"- {table}: {size}")
                    
                    with col2:
                        st.markdown("### Supabase Stats")
                        try:
                            from modules.db.sb.client import get_supabase_client
                            supabase = get_supabase_client()
                            
                            # Get raw listings count from Supabase
                            raw_response = supabase.from_("raw_listings").select("*", count="exact").limit(1).execute()
                            raw_count = raw_response.count
                            
                            # Get enriched listings count from Supabase
                            enriched_response = supabase.from_("enriched_listings").select("*", count="exact").limit(1).execute()
                            enriched_count = enriched_response.count
                            
                            st.markdown(f"**Raw Listings:** {raw_count:,}")
                            st.markdown(f"**Enriched Listings:** {enriched_count:,}")
                            
                        except Exception as supabase_e:
                            st.error(f"Error fetching Supabase stats: {supabase_e}")
                
                except Exception as e:
                    st.error(f"Error fetching database stats: {e}")
                finally:
                    if 'conn' in locals() and conn:
                        conn.close()
        
        with stats_tabs[1]:
            st.subheader("Data Distribution Analysis")
            
            if st.button("Analyze Data Distribution", key="analyze_distribution"):
                try:
                    # Create new connection for this tab
                    from modules.db.ps.database import get_db_connection
                    conn = get_db_connection()
                    
                    with conn.cursor() as cur:
                        # Get distinct search queries
                        cur.execute("SELECT DISTINCT search_query FROM raw_listings ORDER BY search_query")
                        search_queries = [row[0] for row in cur.fetchall()]
                    
                    # Add search query selector
                    selected_query = st.selectbox("Select Search Query", ["All"] + search_queries)
                    
                    # Base query conditions
                    where_clause = "WHERE 1=1" if selected_query == "All" else f"WHERE search_query = '{selected_query}'"
                    
                    with conn.cursor() as cur:
                        # Get website distribution
                        cur.execute(f"""
                            SELECT category, COUNT(*) as count 
                            FROM raw_listings
                            {where_clause}
                            GROUP BY category 
                            ORDER BY count DESC
                        """)
                        website_distribution = cur.fetchall()
                        
                        
                    # Create visualizations
                    col1 = st.columns(1)
                    
                    with col1:
                        # Website distribution chart
                        if website_distribution:
                            website_df = pd.DataFrame(website_distribution, columns=['Website', 'Count'])
                            fig = px.pie(website_df, values='Count', names='Website', title='Listings by Website')
                            st.plotly_chart(fig, use_container_width=True)
                    
                    
                
                except Exception as e:
                    st.error(f"Error analyzing data distribution: {e}")
                finally:
                    if 'conn' in locals() and conn:
                        conn.close()
        
        with stats_tabs[2]:
            st.subheader("Database Performance Metrics")
            
            if st.button("Check Database Performance", key="check_performance"):
                try:
                    from modules.db.ps.database import get_db_connection
                    conn = get_db_connection()
                    
                    with conn.cursor() as cur:
                        # Get database size
                        cur.execute("""
                            SELECT pg_size_pretty(pg_database_size(current_database()))
                        """)
                        db_size = cur.fetchone()[0]
                        
                        # Get index usage statistics
                        cur.execute("""
                            SELECT 
                                relname as table_name,
                                idx_scan as index_scans,
                                seq_scan as sequential_scans,
                                idx_tup_read as tuples_read_via_index,
                                seq_tup_read as tuples_read_via_seq_scan
                            FROM 
                                pg_stat_user_tables
                            WHERE 
                                relname IN ('raw_listings', 'enriched_listings', 'preprocessed_listings')
                            ORDER BY 
                                relname
                        """)
                        index_stats = cur.fetchall()
                        
                        # Get query performance statistics
                        cur.execute("""
                            SELECT 
                                query,
                                calls,
                                round(total_exec_time::numeric, 2) as total_time_ms,
                                round(mean_exec_time::numeric, 2) as mean_time_ms,
                                round(max_exec_time::numeric, 2) as max_time_ms
                            FROM 
                                pg_stat_statements
                            WHERE 
                                query ILIKE '%listings%'
                            ORDER BY 
                                total_exec_time DESC
                            LIMIT 5
                        """)
                        query_stats = cur.fetchall()
                    
                    st.markdown(f"**Database Size:** {db_size}")
                    
                    st.markdown("### Index Usage Statistics")
                    if index_stats:
                        index_df = pd.DataFrame(index_stats, 
                                           columns=['Table', 'Index Scans', 'Sequential Scans', 
                                                    'Tuples Read via Index', 'Tuples Read via Seq Scan'])
                        st.dataframe(index_df)
                    else:
                        st.info("No index usage statistics available")
                    
                    st.markdown("### Top 5 Slowest Queries")
                    if query_stats:
                        query_df = pd.DataFrame(query_stats,
                                           columns=['Query', 'Calls', 'Total Time (ms)', 
                                                    'Mean Time (ms)', 'Max Time (ms)'])
                        st.dataframe(query_df)
                    else:
                        st.info("No query statistics available. You may need to enable pg_stat_statements extension.")
                
                except Exception as e:
                    st.error(f"Error checking database performance: {e}")
                finally:
                    if 'conn' in locals() and conn:
                        conn.close()
        
        with stats_tabs[3]:
            st.subheader("Related Searches Analysis")
            st.markdown("""
            This tab shows the related search terms that have been collected while scraping websites.
            These related searches can provide valuable insights into:
            
            - What other terms are commonly associated with your search queries
            - Market trends and related products
            - Alternative terminology used by sellers
            """)
            
            # Get parameters for filtering
            from modules.db.ps.database import get_db_connection, get_related_searches
            
            try:
                conn = get_db_connection()
                
                col1, col2 = st.columns(2)
                
                with col1:
                    # Get unique search queries
                    with conn.cursor() as cur:
                        cur.execute("SELECT DISTINCT original_search_query FROM related_searches ORDER BY original_search_query")
                        queries = [row[0] for row in cur.fetchall()]
                    
                    if queries:
                        selected_query = st.selectbox("Select Search Query", ["All"] + queries)
                    else:
                        st.info("No related searches found in database yet.")
                        selected_query = "All"
                
                with col2:
                    # Get unique websites
                    with conn.cursor() as cur:
                        cur.execute("SELECT DISTINCT website FROM related_searches ORDER BY website")
                        websites = [row[0] for row in cur.fetchall()]
                    
                    if websites:
                        selected_website = st.selectbox("Select Website", ["All"] + websites)
                    else:
                        selected_website = "All"
                
                # Limit slider
                max_results = st.slider("Maximum Results", min_value=10, max_value=1000, value=100, step=10)
                
                if st.button("Get Related Searches", key="get_related_searches"):
                    # Prepare filter parameters
                    query_param = selected_query if selected_query != "All" else None
                    website_param = selected_website if selected_website != "All" else None
                    
                    # Get related searches from database
                    related_searches_df = get_related_searches(
                        search_query=query_param,
                        website=website_param,
                        limit=max_results
                    )
                    
                    if related_searches_df.empty:
                        st.info("No related searches found with the selected filters.")
                    else:
                        st.success(f"Found {len(related_searches_df)} related search terms")
                        
                        # Show data table
                        st.dataframe(related_searches_df)
                        
                        # Visualize top terms if there are enough
                        if len(related_searches_df) >= 5:
                            st.subheader("Top Related Terms by Search Count")
                            
                            # Prepare data for visualization
                            plot_df = related_searches_df.sort_values("search_count", ascending=False).head(20)
                            
                            fig = px.bar(
                                plot_df,
                                x="related_term",
                                y="search_count",
                                color="website" if "website" in plot_df.columns and len(plot_df["website"].unique()) > 1 else None,
                                title=f"Top 20 Related Terms for {selected_query if selected_query != 'All' else 'All Queries'}"
                            )
                            
                            fig.update_layout(
                                xaxis_title="Related Term",
                                yaxis_title="Occurrence Count",
                                xaxis={'categoryorder':'total descending'}
                            )
                            
                            st.plotly_chart(fig, use_container_width=True)
                            
                            # Option to download the data
                            csv = related_searches_df.to_csv(index=False)
                            st.download_button(
                                label="Download CSV",
                                data=csv,
                                file_name=f"related_searches_{selected_query if selected_query != 'All' else 'all'}_{selected_website if selected_website != 'All' else 'all'}.csv",
                                mime="text/csv"
                            )
            
            except Exception as e:
                st.error(f"Error accessing related searches data: {e}")
            finally:
                if 'conn' in locals() and conn:
                    conn.close()

        with stats_tabs[4]:
            st.subheader("Search Listings")
            st.markdown("""
            Search for listings using similarity search on titles and apply filters.
            Results can be viewed in this tab or exported for further analysis.
            """)
            
            # Database source selection
            db_source = st.radio("Select Database Source", ["PostgreSQL", "Supabase"], horizontal=True)
            
            # Create columns for search input
            col1, col2 = st.columns([3, 1])
            
            with col1:
                search_term = st.text_input("Search Term (for title similarity)", 
                                           placeholder="Enter search text for title similarity search...")
            
            with col2:
                limit = st.number_input("Max Results", min_value=1, max_value=1000, value=100)
            
            # Create collapsible section for filters
            with st.expander("Search Filters", expanded=False):
                filter_col1, filter_col2 = st.columns(2)
                
                with filter_col1:
                    # Date range filter
                    st.subheader("Date Filter")
                    date_filter_enabled = st.checkbox("Filter by Date")
                    
                    if date_filter_enabled:
                        date_range = st.date_input("Date Range", 
                                                  value=(datetime.datetime.now() - datetime.timedelta(days=90), 
                                                        datetime.datetime.now()),
                                                  format="YYYY-MM-DD")
                    else:
                        date_range = None
                    
                    # Search query filter
                    st.subheader("Search Query")
                    search_query_filter_enabled = st.checkbox("Filter by Search Query")
                    
                    if search_query_filter_enabled:
                        # Get unique search queries
                        try:
                            if db_source == "PostgreSQL":
                                from modules.db.ps.database import get_db_connection
                                conn = get_db_connection()
                                with conn.cursor() as cur:
                                    cur.execute("SELECT DISTINCT search_query FROM raw_listings ORDER BY search_query")
                                    search_queries = [row[0] for row in cur.fetchall()]
                                conn.close()
                            else:  # Supabase
                                from modules.db.sb.client import get_supabase_client
                                supabase = get_supabase_client()
                                response = supabase.from_("raw_listings").select("search_query").limit(1000).execute()
                                search_queries = list(set([item['search_query'] for item in response.data if 'search_query' in item]))
                                search_queries.sort()
                            
                            selected_search_query = st.selectbox("Select Search Query", search_queries)
                        except Exception as e:
                            st.error(f"Error loading search queries: {e}")
                            selected_search_query = None
                    else:
                        selected_search_query = None
                    
                    # Category filter
                    st.subheader("Category")
                    category_filter_enabled = st.checkbox("Filter by Category")
                    
                    if category_filter_enabled:
                        # Get unique categories
                        try:
                            if db_source == "PostgreSQL":
                                from modules.db.ps.database import get_db_connection
                                conn = get_db_connection()
                                with conn.cursor() as cur:
                                    cur.execute("SELECT DISTINCT category FROM raw_listings ORDER BY category")
                                    categories = [row[0] for row in cur.fetchall()]
                                conn.close()
                            else:  # Supabase
                                from modules.db.sb.client import get_supabase_client
                                supabase = get_supabase_client()
                                response = supabase.from_("raw_listings").select("category").limit(1000).execute()
                                categories = list(set([item['category'] for item in response.data if 'category' in item]))
                                categories.sort()
                            
                            selected_category = st.selectbox("Select Category", categories)
                        except Exception as e:
                            st.error(f"Error loading categories: {e}")
                            selected_category = None
                    else:
                        selected_category = None
                
                with filter_col2:
                    # Usage filter
                    st.subheader("Usage")
                    usage_filter_enabled = st.checkbox("Filter by Usage")
                    
                    if usage_filter_enabled:
                        # Get unique usage values or provide common ones
                        usage_options = ["New", "Used", "Unknown"]
                        selected_usage = st.selectbox("Select Usage", usage_options)
                    else:
                        selected_usage = None
                    
                    # Price range filter
                    st.subheader("Price Range")
                    price_filter_enabled = st.checkbox("Filter by Price")
                    
                    if price_filter_enabled:
                        price_min = st.number_input("Min Price", min_value=0, value=0)
                        price_max = st.number_input("Max Price", min_value=0, value=10000)
                        price_range = (price_min, price_max)
                    else:
                        price_range = None
                    
                    # Website filter
                    st.subheader("Website")
                    website_filter_enabled = st.checkbox("Filter by Website")
                    
                    if website_filter_enabled:
                        # Get unique websites
                        try:
                            if db_source == "PostgreSQL":
                                from modules.db.ps.database import get_db_connection
                                conn = get_db_connection()
                                with conn.cursor() as cur:
                                    cur.execute("SELECT DISTINCT website FROM raw_listings ORDER BY website")
                                    websites = [row[0] for row in cur.fetchall()]
                                conn.close()
                            else:  # Supabase
                                from modules.db.sb.client import get_supabase_client
                                supabase = get_supabase_client()
                                response = supabase.from_("raw_listings").select("website").limit(1000).execute()
                                websites = list(set([item['website'] for item in response.data if 'website' in item]))
                                websites.sort()
                            
                            selected_website = st.selectbox("Select Website", websites)
                        except Exception as e:
                            st.error(f"Error loading websites: {e}")
                            selected_website = None
                    else:
                        selected_website = None
                    
                    # Company filter
                    st.subheader("Company/Location")
                    company_filter_enabled = st.checkbox("Filter by Company/Location")
                    
                    if company_filter_enabled:
                        # Text input for company/location
                        selected_company = st.text_input("Company or Location", 
                                                       placeholder="Enter company or location...")
                    else:
                        selected_company = None
            
            # Build search fields dictionary
            search_fields = {}
            if date_filter_enabled and date_range:
                search_fields['date_range'] = date_range
            if search_query_filter_enabled and selected_search_query:
                search_fields['search_query'] = selected_search_query
            if category_filter_enabled and selected_category:
                search_fields['category'] = selected_category
            if usage_filter_enabled and selected_usage:
                search_fields['usage'] = selected_usage
            if price_filter_enabled and price_range:
                search_fields['price_range'] = price_range
            if website_filter_enabled and selected_website:
                search_fields['website'] = selected_website
            if company_filter_enabled and selected_company:
                search_fields['company'] = selected_company
            
            # Search button
            if st.button("Search Listings", key="search_listings"):
                try:
                    if not search_term and not search_fields:
                        st.warning("Please enter a search term or select at least one filter.")
                    else:
                        results_df = None
                        
                        if db_source == "PostgreSQL":
                            from modules.db.ps.database import search_listings_by_similarity
                            results_df = search_listings_by_similarity(search_term, search_fields, limit)
                        else:  # Supabase
                            from modules.db.sb.client import search_listings_in_supabase
                            results_df = search_listings_in_supabase(search_term, search_fields, limit)
                        
                        if results_df is not None and not results_df.empty:
                            st.session_state['search_results'] = results_df
                            st.success(f"Found {len(results_df)} matching listings.")
                            
                            # Display results in a table
                            st.dataframe(results_df, use_container_width=True)
                            
                            # Export options
                            st.download_button(
                                label="Download Results as CSV",
                                data=results_df.to_csv(index=False).encode('utf-8'),
                                file_name=f"search_results_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                mime="text/csv"
                            )
                        else:
                            st.info("No matching listings found.")
                
                except Exception as e:
                    st.error(f"Error searching listings: {e}")
                    import traceback
                    st.error(traceback.format_exc())
    # Page 2: Enrich Data
    elif st.session_state.page == "enrich_data":
        st.header("Transform Data")
        
        # Create tabs for different enrichment functions
        enrich_tabs = st.tabs(["Data Preprocessing", "Run Enrichment", "Batch Processing", "Embeddings Management"])
        
        with enrich_tabs[0]:
            st.subheader("Data Preprocessing")
            st.markdown("""
            Preprocessing converts raw data into a standardized format by:
            - Cleaning price strings to numeric values
            - Parsing various date formats
            - Preparing data for enrichment
            """)
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Button to run preprocessing on new raw listings
                st.subheader("Process Raw Listings")
                
                # Add filters for preprocessing
                from modules.db.ps.database import get_db_connection
                
                # Get available search queries, categories, and sample titles
                search_queries = ["All"]
                categories = ["All"]
                sample_titles = ["All"]
                
                try:
                    conn = get_db_connection()
                    with conn.cursor() as cur:
                        # Get distinct search queries
                        cur.execute("""
                            SELECT DISTINCT search_query FROM raw_listings 
                            WHERE search_query IS NOT NULL 
                            ORDER BY search_query
                        """)
                        search_queries.extend([row[0] for row in cur.fetchall()])
                        
                        # Get distinct categories
                        cur.execute("""
                            SELECT DISTINCT category FROM raw_listings 
                            WHERE category IS NOT NULL 
                            ORDER BY category
                        """)
                        categories.extend([row[0] for row in cur.fetchall()])
                        
                        # Get sample titles (limited to 20)
                        cur.execute("""
                            SELECT DISTINCT title FROM raw_listings 
                            WHERE title IS NOT NULL 
                            LIMIT 20
                        """)
                        sample_titles.extend([row[0] for row in cur.fetchall()])
                except Exception as e:
                    st.error(f"Error fetching filter options: {e}")
                finally:
                    if 'conn' in locals() and conn:
                        conn.close()
                
                # Filter options
                filter_type = st.radio("Filter by:", ["Search Query", "Category", "Title"], horizontal=True)
                
                if filter_type == "Search Query":
                    selected_filter = st.selectbox("Select Search Query:", search_queries)
                    filter_clause = "" if selected_filter == "All" else f"r.search_query = '{selected_filter}'"
                elif filter_type == "Category":
                    selected_filter = st.selectbox("Select Category:", categories)
                    filter_clause = "" if selected_filter == "All" else f"r.category = '{selected_filter}'"
                else:  # Title
                    selected_filter = st.selectbox("Select Title:", sample_titles)
                    filter_clause = "" if selected_filter == "All" else f"r.title = '{selected_filter}'"
                
                # Allow user to specify batch size and limit
                batch_size = st.slider("Batch Size", min_value=10, max_value=500, value=100, step=10)
                limit = st.slider("Max Listings to Process", min_value=100, max_value=5000, value=1000, step=100)
                
                if st.button("Process Selected Raw Listings", key="process_raw"):
                    with st.spinner("Processing raw listings from database..."):
                        try:
                            from modules.preprocessing_pipeline import process_raw_listings_batch
                            
                            try:
                                # Try with filter_clause parameter first
                                processed, errors = process_raw_listings_batch(
                                    batch_size=batch_size, 
                                    limit=limit,
                                    filter_clause=filter_clause
                                )
                            except TypeError as e:
                                # If that fails because filter_clause isn't accepted, try without it
                                if "unexpected keyword argument 'filter_clause'" in str(e):
                                    st.warning("Using older version of preprocessing pipeline that doesn't support filtering.")
                                    processed, errors = process_raw_listings_batch(
                                        batch_size=batch_size, 
                                        limit=limit
                                    )
                                else:
                                    raise
                            
                            if processed > 0:
                                st.success(f"Successfully processed {processed} raw listings (with {errors} errors)")
                            else:
                                st.info("No new raw listings to process")
                                
                        except Exception as e:
                            st.error(f"Error during preprocessing: {e}")
            
            with col2:
                # Show preprocessing statistics
                st.subheader("Preprocessing Stats")
                
                if st.button("Refresh Stats", key="refresh_preprocess_stats"):
                    from modules.db.ps.database import get_db_connection
                    
                    try:
                        conn = get_db_connection()
                        with conn.cursor() as cur:
                            # Get counts of processed and unprocessed listings
                            cur.execute("""
                                SELECT 
                                    (SELECT COUNT(*) FROM raw_listings) as total_raw,
                                    (SELECT COUNT(*) FROM preprocessed_listings) as total_preprocessed,
                                    (SELECT COUNT(*) FROM raw_listings r 
                                     LEFT JOIN preprocessed_listings p ON r.id = p.raw_listing_id
                                     WHERE p.id IS NULL) as unprocessed
                            """)
                            
                            stats = cur.fetchone()
                            
                            if stats:
                                total_raw, total_preprocessed, unprocessed = stats
                                
                                st.metric("Total Raw Listings", f"{total_raw:,}")
                                st.metric("Total Preprocessed", f"{total_preprocessed:,}")
                                st.metric("Pending Preprocessing", f"{unprocessed:,}")
                                
                                if total_raw > 0:
                                    progress = (total_preprocessed / total_raw) * 100
                                    st.progress(min(progress / 100, 1.0), text=f"{progress:.1f}% Processed")
                    except Exception as e:
                        st.error(f"Error fetching preprocessing stats: {e}")
                    finally:
                        if 'conn' in locals() and conn:
                            conn.close()
                
        with enrich_tabs[1]:
            st.subheader("Data Enrichment")
            st.markdown("""
            Enrichment adds value to preprocessed data by:
            - Classifying the brand and model using LLM or embeddings
            - Using LLM for the first 20+ samples of a search query
            - Using embedding similarity for subsequent classifications (ranky method)
            """)
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Button to run enrichment on preprocessed listings
                if st.button("Enrich Preprocessed Listings", key="enrich_data"):
                    with st.spinner("Enriching preprocessed listings..."):
                        try:
                            import asyncio
                            from modules.enrichment_pipeline import process_preprocessed_listings_batch
                            
                            # Allow user to specify batch size and limit
                            batch_size = st.slider("Batch Size", min_value=5, max_value=50, value=10, step=5)
                            limit = st.slider("Max Listings to Enrich", min_value=10, max_value=500, value=50, step=10)
                            
                            # Run the async function
                            processed, errors = asyncio.run(process_preprocessed_listings_batch(batch_size=batch_size, limit=limit))
                            
                            if processed > 0:
                                st.success(f"Successfully enriched {processed} preprocessed listings (with {errors} errors)")
                            else:
                                st.info("No new preprocessed listings to enrich")
                                
                        except Exception as e:
                            st.error(f"Error during enrichment: {e}")
                            
            with col2:
                # Show enrichment statistics
                st.subheader("Enrichment Stats")
                
                if st.button("Refresh Stats", key="refresh_enrich_stats"):
                    from modules.db.ps.database import get_db_connection
                    
                    try:
                        conn = get_db_connection()
                        with conn.cursor() as cur:
                            # Get counts of preprocessed and enriched listings
                            cur.execute("""
                                SELECT 
                                    (SELECT COUNT(*) FROM preprocessed_listings) as total_preprocessed,
                                    (SELECT COUNT(*) FROM enriched_listings) as total_enriched,
                                    (SELECT COUNT(*) FROM preprocessed_listings p 
                                     LEFT JOIN enriched_listings e ON p.id = e.preprocessed_listing_id
                                     WHERE e.id IS NULL
                                     AND p.processing_status = 'processed') as unenriched
                            """)
                            
                            stats = cur.fetchone()
                            
                            if stats:
                                total_preprocessed, total_enriched, unenriched = stats
                                
                                st.metric("Total Preprocessed", f"{total_preprocessed:,}")
                                st.metric("Total Enriched", f"{total_enriched:,}")
                                st.metric("Pending Enrichment", f"{unenriched:,}")
                                
                                if total_preprocessed > 0:
                                    progress = (total_enriched / total_preprocessed) * 100
                                    st.progress(min(progress / 100, 1.0), text=f"{progress:.1f}% Enriched")
                                    
                            # Get enrichment method stats
                            cur.execute("""
                                SELECT COUNT(*) FROM embeddings
                                WHERE table_source = 'preprocessed_listings'
                            """)
                            
                            embedding_count = cur.fetchone()[0]
                            st.metric("Stored Embeddings", f"{embedding_count:,}")
                    except Exception as e:
                        st.error(f"Error fetching enrichment stats: {e}")
                    finally:
                        if 'conn' in locals() and conn:
                            conn.close()
        
        with enrich_tabs[2]:
            st.subheader("Batch Processing")
            st.markdown("""
            Process multiple steps at once for efficiency:
            1. Preprocess raw listings
            2. Enrich preprocessed listings
            3. Update embedding centroids for classification
            """)
            
            if st.button("Run Complete Pipeline", key="run_pipeline"):
                with st.spinner("Running complete data pipeline..."):
                    try:
                        import asyncio
                        from modules.preprocessing_pipeline import process_raw_listings_batch
                        from modules.enrichment_pipeline import process_preprocessed_listings_batch, update_centroids_for_search_query
                        
                        # Step 1: Preprocess raw listings
                        st.text("Step 1: Preprocessing raw listings...")
                        processed_raw, errors_raw = process_raw_listings_batch(batch_size=100, limit=500)
                        st.text(f"Preprocessed {processed_raw} raw listings with {errors_raw} errors")
                        
                        # Step 2: Enrich preprocessed listings
                        st.text("Step 2: Enriching preprocessed listings...")
                        processed_enriched, errors_enriched = asyncio.run(process_preprocessed_listings_batch(batch_size=10, limit=100))
                        st.text(f"Enriched {processed_enriched} preprocessed listings with {errors_enriched} errors")
                        
                        # Step 3: Update centroids for active search queries
                        st.text("Step 3: Updating centroids for active search queries...")
                        
                        # Get active search queries
                        from modules.db.ps.database import get_db_connection
                        conn = get_db_connection()
                        search_queries = []
                        
                        try:
                            with conn.cursor() as cur:
                                cur.execute("SELECT DISTINCT search_query FROM enriched_listings WHERE search_query IS NOT NULL LIMIT 5")
                                search_queries = [row[0] for row in cur.fetchall()]
                        finally:
                            conn.close()
                        
                        total_centroids = 0
                        for query in search_queries:
                            if query:
                                brand_centroids = asyncio.run(update_centroids_for_search_query(query, "brand"))
                                model_centroids = asyncio.run(update_centroids_for_search_query(query, "model"))
                                total_centroids += brand_centroids + model_centroids
                                st.text(f"Updated {brand_centroids} brand and {model_centroids} model centroids for '{query}'")
                        
                        st.success(f"""
                        Pipeline completed:
                        - Preprocessed: {processed_raw} listings ({errors_raw} errors)
                        - Enriched: {processed_enriched} listings ({errors_enriched} errors)
                        - Updated: {total_centroids} centroids
                        """)
                        
                    except Exception as e:
                        st.error(f"Error running pipeline: {e}")
        
        with enrich_tabs[3]:
            st.subheader("Embeddings Management")
            st.markdown("""
            Manage embeddings and centroids for the classification system:
            - Generate missing embeddings
            - Update centroids for specific search queries
            - View embedding statistics
            """)
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Form to update centroids for a specific search query
                with st.form("update_centroids_form"):
                    st.subheader("Update Centroids")
                    
                    # Get available search queries
                    from modules.db.ps.database import get_db_connection
                    search_queries = []
                    
                    try:
                        conn = get_db_connection()
                        with conn.cursor() as cur:
                            cur.execute("SELECT DISTINCT search_query FROM enriched_listings WHERE search_query IS NOT NULL ORDER BY search_query")
                            search_queries = [row[0] for row in cur.fetchall()]
                    except Exception as e:
                        st.error(f"Error fetching search queries: {e}")
                    finally:
                        if 'conn' in locals() and conn:
                            conn.close()
                    
                    selected_query = st.selectbox("Select Search Query", search_queries if search_queries else ["No queries available"])
                    category_type = st.radio("Category Type", ["Brand", "Model"], horizontal=True)
                    
                    submit_centroids = st.form_submit_button("Update Centroids")
                    
                    if submit_centroids and selected_query and selected_query != "No queries available":
                        with st.spinner(f"Updating {category_type.lower()} centroids for '{selected_query}'..."):
                            try:
                                import asyncio
                                from modules.enrichment_pipeline import update_centroids_for_search_query
                                
                                updated = asyncio.run(update_centroids_for_search_query(selected_query, category_type.lower()))
                                
                                if updated > 0:
                                    st.success(f"Updated {updated} {category_type.lower()} centroids for '{selected_query}'")
                                else:
                                    st.info(f"No {category_type.lower()} centroids to update for '{selected_query}'")
                                    
                            except Exception as e:
                                st.error(f"Error updating centroids: {e}")
            
            with col2:
                # Show embedding and centroid statistics
                st.subheader("Embedding Statistics")
                
                if st.button("Refresh Embedding Stats", key="refresh_embedding_stats"):
                    from modules.db.ps.database import get_db_connection
                    
                    try:
                        conn = get_db_connection()
                        with conn.cursor() as cur:
                            # Get embedding counts by table source
                            cur.execute("""
                                SELECT table_source, COUNT(*) 
                                FROM embeddings 
                                GROUP BY table_source
                            """)
                            
                            embedding_stats = cur.fetchall()
                            
                            if embedding_stats:
                                st.markdown("**Embeddings by Source:**")
                                for source, count in embedding_stats:
                                    st.metric(source, f"{count:,}")
                            
                            # Get centroid counts by category type
                            cur.execute("""
                                SELECT category_type, COUNT(*) 
                                FROM centroids 
                                GROUP BY category_type
                            """)
                            
                            centroid_stats = cur.fetchall()
                            
                            if centroid_stats:
                                st.markdown("**Centroids by Category Type:**")
                                for category_type, count in centroid_stats:
                                    st.metric(category_type, f"{count:,}")
                                    
                            # Get top 5 search queries by centroid count
                            cur.execute("""
                                SELECT search_query, COUNT(*) 
                                FROM centroids 
                                WHERE search_query IS NOT NULL
                                GROUP BY search_query 
                                ORDER BY COUNT(*) DESC
                                LIMIT 5
                            """)
                            
                            top_queries = cur.fetchall()
                            
                            if top_queries:
                                st.markdown("**Top Search Queries by Centroid Count:**")
                                for query, count in top_queries:
                                    st.text(f"{query}: {count} centroids")
                    except Exception as e:
                        st.error(f"Error fetching embedding stats: {e}")
                    finally:
                        if 'conn' in locals() and conn:
                            conn.close()
    # Page 3: Analysis
    elif st.session_state.page == "analytics":
        st.header("📈 Analytics")

        # Retrieve the enriched DataFrame or preprocessed DataFrame if enriched is not available
        df_enriched = st.session_state.get("enriched_df", None)
        df_preprocessed = st.session_state.get("preprocessed_df", None)
        
        if df_enriched is None and df_preprocessed is None:
            st.warning("Please load and preprocess data first before analysis.")
        else:
            # Use enriched data if available, otherwise use preprocessed data
            df_analysis = df_enriched if df_enriched is not None else df_preprocessed
            is_enriched = df_enriched is not None

            # Sidebar Filters
            st.sidebar.header("🔍 Filter Options")
            
            # Brand Selection (if enriched data available)
            if is_enriched and 'Brand' in df_analysis.columns:
                available_brands = sorted(df_analysis['Brand'].dropna().unique())
                selected_brands = st.sidebar.multiselect(
                    "🏢 Select Brands:",
                    options=available_brands,
                    default=available_brands
                )
            else:
                selected_brands = None
            
            # Usage Selection
            available_usages = sorted(df_analysis['Usage'].dropna().unique())
            selected_usages = st.sidebar.multiselect(
                "🔧 Select Usage Categories:",
                options=available_usages,
                default=available_usages
            )
            
            # Date Range Selection
            selected_date_range = st.sidebar.date_input(
                "📅 Select Date Range:",
            )
            
            # Price Range Selection
            price_col = 'CleanedPrice' if 'CleanedPrice' in df_analysis.columns else 'Price'
            numeric_prices = pd.to_numeric(df_analysis[price_col], errors='coerce')
            valid_prices = numeric_prices.dropna()

            if valid_prices.empty:
                st.warning("No valid numeric prices available for filtering.")
                selected_price_range = (0.0, 0.0)
            else:
                min_price = float(valid_prices.min())
                max_price = float(valid_prices.max())
                selected_price_range = st.sidebar.slider(
                    "💰 Price Range",
                    min_value=min_price,
                    max_value=max_price,
                    value=(min_price, max_price)
                )

            # Apply filters
            filtered_df = df_analysis.copy()
            
            # Apply brand filter
            if selected_brands:
                filtered_df = filtered_df[filtered_df['Brand'].isin(selected_brands)]
            
            # Apply usage filter
            filtered_df = filtered_df[filtered_df['Usage'].isin(selected_usages)]
            
            # Apply price filter
            filtered_df = filtered_df[
                (pd.to_numeric(filtered_df[price_col], errors='coerce') >= selected_price_range[0]) &
                (pd.to_numeric(filtered_df[price_col], errors='coerce') <= selected_price_range[1])
            ]
            
            # Apply date filter if provided
            if isinstance(selected_date_range, tuple) and len(selected_date_range) == 2:
                start_date, end_date = selected_date_range
                filtered_df = filtered_df[
                    (pd.to_datetime(filtered_df['Date']) >= pd.to_datetime(start_date)) &
                    (pd.to_datetime(filtered_df['Date']) <= pd.to_datetime(end_date))
                ]

            st.markdown(f"### 📊 Selected Data: {len(filtered_df)} Listings")
            st.dataframe(filtered_df.head(10))

            # Analysis selection
            analysis_type = st.sidebar.radio(
                "Select Analysis Type",
                ["📈 Basic Analysis", "📝 Sentiment Analysis", "🏆 Best Deal", "🔍 Market Insights", "📊 Price Trends"] if is_enriched else ["📈 Basic Analysis"]
            )

            # Basic Analysis
            if analysis_type == "📈 Basic Analysis":
                st.subheader("📈 Basic Analysis")
                
                # Price Statistics
                st.write("### 💰 Price Analysis")
                numeric_prices = pd.to_numeric(filtered_df[price_col], errors='coerce')
                price_stats = numeric_prices.describe()
                st.write(price_stats)
                
                # Price Distribution Plot
                fig_price = px.histogram(
                    filtered_df[pd.to_numeric(filtered_df[price_col], errors='coerce').notna()], 
                    x=price_col, 
                    title="Price Distribution"
                )
                st.plotly_chart(fig_price)
                
                # Usage Distribution
                st.write("### 🔧 Usage Distribution")
                usage_counts = filtered_df['Usage'].value_counts()
                fig_usage = px.pie(
                    values=usage_counts.values,
                    names=usage_counts.index,
                    title="Usage Distribution"
                )
                st.plotly_chart(fig_usage)
                
                # Price Trends Over Time
                st.write("### 📅 Price Trends Over Time")
                trend_df = filtered_df[pd.to_numeric(filtered_df[price_col], errors='coerce').notna()]
                fig_trend = px.scatter(
                    trend_df,
                    x='Date',
                    y=price_col,
                    title="Price Trends Over Time"
                )
                st.plotly_chart(fig_trend)

                # Add download report button
                if st.button("📥 Generate Analysis Report"):
                    with st.spinner("Generating comprehensive report..."):
                        download_analysis_report(filtered_df)

            # Sentiment Analysis
            elif analysis_type == "📝 Sentiment Analysis" and is_enriched:
                st.subheader("📝 Sentiment Analysis")
                if st.button("Analyze Sentiment"):
                    with st.spinner("🧠 Performing sentiment analysis..."):
                        try:
                            enriched_df = asyncio.run(sentiment_analysis_agent(filtered_df))
                            st.success("✅ Sentiment analysis complete.")
                            st.dataframe(enriched_df[['Title', 'Sentiment', 'SentimentScore']].head(10))
                        except Exception as e:
                            st.error(f"❌ Error during sentiment analysis: {e}")

            # Best Deal
            elif analysis_type == "🏆 Best Deal" and is_enriched:
                st.subheader("🏆 Best Deal Recommendation")
                if st.button("Get Best Deal"):
                    with st.spinner("🔍 Identifying the best deal..."):
                        try:
                            best_deal = asyncio.run(best_deal_recommendation(filtered_df))
                            st.info(f"**Best Deal Recommendation:**\n{best_deal}")
                        except Exception as e:
                            st.error(f"❌ Error during best deal recommendation: {e}")

            # Market Insights
            elif analysis_type == "🔍 Market Insights" and is_enriched:
                st.subheader("🔍 Market Insights")
                if st.button("Generate Market Insights"):
                    with st.spinner("🧠 Generating market insights..."):
                        try:
                            insights = asyncio.run(market_insights_analysis(filtered_df))
                            st.session_state["insights"] = insights
                            st.info(f"**Market Insights:**\n{insights}")
                        except Exception as e:
                            st.error(f"❌ Error during market insights analysis: {e}")

            # Price Trend Analysis
            elif analysis_type == "📊 Price Trends" and is_enriched:
                st.subheader("📊 Price Trend Analysis")
                try:
                    trend_analysis = asyncio.run(price_trend_analysis(filtered_df))
                    st.markdown(trend_analysis)
                    
                    if 'Date' in filtered_df.columns and 'CleanedPrice' in filtered_df.columns:
                        filtered_df['Date'] = pd.to_datetime(filtered_df['Date'])
                        filtered_df = filtered_df.sort_values('Date')
                        fig = px.line(
                            filtered_df,
                            x='Date',
                            y='CleanedPrice',
                            title='Price Trends Over Time'
                        )
                        st.plotly_chart(fig)
                except Exception as e:
                    st.error(f"❌ Error during price trend analysis: {e}")

    # Page 4: Email Reports
    elif st.session_state.page == "reports":
        st.header("📧 Reports")
        
        insights = st.session_state.get("insights", "")
        if not insights:
            st.warning("⚠️ No insights available to send.")
            return

        # Create two columns for email and WhatsApp sections
        col1, col2 = st.columns(2)

        # Email Section
        with col1:
            st.subheader("📧 Send via Email")
            with st.form("email_form"):
                recipient_email = st.text_input("Recipient Email", value="david.kakanis@hotmail.com")
                recipient_name = st.text_input("Recipient Name", value="John Doe")
                submit_email = st.form_submit_button("Send Email")

                if submit_email:
                    if not recipient_email:
                        st.error("❌ Please provide an email address.")
                    elif not is_valid_email(recipient_email):
                        st.error("❌ Please provide a valid email address.")
                    else:
                        with st.spinner("Generating and sending report..."):
                            send_email_insights(insights, recipient_email, recipient_name)
        # Slack Section
        with col2:
            st.subheader("💬 Send via Slack")
            with st.form("slack_form"):
                slack_channel = st.text_input(
                    "Slack Channel",
                    value="#market-insights",
                    help="Enter the Slack channel name (e.g., #general)"
                )
                submit_slack = st.form_submit_button("Send to Slack")

                if submit_slack:
                    if not slack_channel or not slack_channel.startswith("#"):
                        st.error("❌ Please provide a valid Slack channel starting with #.")
                    else:
                        with st.spinner("Sending Slack message..."):
                            # Import the send_slack_message function from modules.slack
                            from modules.slack import send_slack_message
                            send_slack_message(insights, slack_channel)
                            st.success("✅ Slack message sent successfully!")

    # Dashboard section
    elif st.session_state.page == "dashboard":
        st.header("📊 Dashboard")
        df_dash = st.session_state.get("enriched_df", pd.DataFrame())
        if df_dash.empty:
            st.warning("⚠️ No data to analyze. Please enrich data first.")
        else:
            # Create metrics for the top of the dashboard
            metric_col1, metric_col2, metric_col3 = st.columns(3)
            with metric_col1:
                try:
                    avg_price = round(df_dash['CleanedPrice'].mean(), 2)
                    st.metric(label="Average Price", value=f"€{avg_price}")
                except:
                    st.metric(label="Average Price", value="N/A")
            
            with metric_col2:
                try:
                    median_price = round(df_dash['CleanedPrice'].median(), 2)
                    st.metric(label="Median Price", value=f"€{median_price}")
                except:
                    st.metric(label="Median Price", value="N/A")
            
            with metric_col3:
                listing_count = len(df_dash)
                st.metric(label="Total Listings", value=listing_count)
            
            # Create visualizations
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("💰 Price Distribution")
                try:
                    fig_price = px.histogram(df_dash, x='CleanedPrice', title="Price Distribution")
                    st.plotly_chart(fig_price, use_container_width=True)
                except Exception as e:
                    st.error(f"Could not create price distribution chart: {e}")
            
            with col2:
                st.subheader("🔧 Usage Distribution")
                try:
                    usage_counts = df_dash['Usage'].value_counts()
                    fig_usage = px.pie(values=usage_counts.values, names=usage_counts.index, title="Usage Categories")
                    st.plotly_chart(fig_usage, use_container_width=True)
                except Exception as e:
                    st.error(f"Could not create usage distribution chart: {e}")
            
            # Price trends over time
            st.subheader("📈 Price Trends Over Time")
            try:
                df_dash['Date'] = pd.to_datetime(df_dash['Date'])
                df_trend = df_dash.groupby(df_dash['Date'].dt.date)['CleanedPrice'].mean().reset_index()
                fig_trend = px.line(
                    df_trend,
                    x='Date',
                    y='CleanedPrice',
                    title="Average Price Over Time"
                )
                st.plotly_chart(fig_trend, use_container_width=True)
            except Exception as e:
                st.error(f"Could not create price trends chart: {e}")
            
            # Display the dataframe with the most important columns
            st.subheader("📊 Recent Listings")
            try:
                display_cols = ['Title', 'CleanedPrice', 'Usage', 'Date']
                if 'Brand' in df_dash.columns:
                    display_cols.insert(2, 'Brand')
                st.dataframe(df_dash[display_cols].head(10), use_container_width=True)
            except Exception as e:
                st.error(f"Could not display listings: {e}")
            
            # Database Operations section
            st.header("Database Operations Dashboard")
            
            # Add a sub-tab for database operations
            db_tabs = st.tabs(["Database Status", "Recent Listings", "Storage Stats"])
            
            with db_tabs[0]:
                st.subheader("Database Connection Status")
                
                if st.button("Check Database Connections"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("### PostgreSQL Status")
                        try:
                            # Try to create tables which will test the connection
                            create_postgres_tables()
                            st.success("✅ Connected to PostgreSQL")
                        except Exception as e:
                            st.error(f"❌ PostgreSQL connection failed: {e}")
                    
                    with col2:
                        st.markdown("### Supabase Status")
                        try:
                            # Import Supabase client
                            from modules.db.sb.client import get_supabase_client
                            client = get_supabase_client()
                            st.success("✅ Connected to Supabase")
                        except Exception as e:
                            st.error(f"❌ Supabase connection failed: {e}")
            
            with db_tabs[1]:
                st.subheader("Recent Listings from Database")
                
                # Add selection for number of listings
                num_listings = st.slider("Number of listings to display", 5, 50, 10)
                
                if st.button("Fetch Recent Listings"):
                    try:
                        st.info(f"Fetching {num_listings} recent listings from database...")
                        listings_df = get_latest_listings(limit=num_listings)
                        
                        if listings_df is not None and not listings_df.empty:
                            st.success(f"Found {len(listings_df)} recent listings")
                            st.dataframe(listings_df)
                            
                            # Add a quick visualization of the data
                            st.subheader("Price Distribution")
                            try:
                                # Clean the price column for visualization
                                listings_df['cleaned_price'] = listings_df['price'].apply(lambda x: clean_price(x))
                                listings_df['cleaned_price'] = pd.to_numeric(listings_df['cleaned_price'], errors='coerce')
                                
                                # Create a price histogram
                                fig = px.histogram(
                                    listings_df, 
                                    x='cleaned_price',
                                    nbins=20,
                                    title='Price Distribution of Recent Listings'
                                )
                                st.plotly_chart(fig, use_container_width=True)
                            except Exception as viz_e:
                                st.error(f"Error creating visualization: {viz_e}")
                        else:
                            st.warning("No recent listings found in the database")
                    except Exception as e:
                        st.error(f"Error fetching recent listings: {e}")
            
            
            
            # Background Analysis section
            st.header("⚙️ Background Analysis")
            st.info("Background analysis functionality is currently under development.")

    # Comprehensive Analysis page
    elif st.session_state.page == "comprehensive":
        st.header("🔍 Comprehensive Analysis")

        # Sub-tabs for different comprehensive analyses
        comprehensive_tabs = st.tabs([
            "📈 Market Demand Forecasting",
            "🔄 Detailed Comparative Analysis",
            "🤖 Agent-Based Analysis"
        ])

        # Sub-tab 1: Market Demand Forecasting
        with comprehensive_tabs[0]:
            st.subheader("📈 Market Demand Forecasting")
            if st.button("Forecast Market Demand"):
                df_for_forecasting = st.session_state.get("enriched_df", pd.DataFrame())
                if df_for_forecasting.empty:
                    st.warning("⚠️ No data to analyze. Please enrich data first.")
                else:
                    with st.spinner("📈 Forecasting market demand..."):
                        try:
                            forecasting = asyncio.run(market_demand_forecasting(df_for_forecasting))
                            st.info(f"**Market Demand Forecasting:**\n{forecasting}")
                        except Exception as e:
                            st.error(f"❌ Error during market demand forecasting: {e}")

        # Sub-tab 2: Detailed Comparative Analysis
        with comprehensive_tabs[1]:
            st.subheader("🔄 Detailed Comparative Analysis")
            if st.button("Compare Listings"):
                df_for_comparison = st.session_state.get("enriched_df", pd.DataFrame())
                if df_for_comparison.empty:
                    st.warning("⚠️ No data to analyze. Please enrich data first.")
                else:
                    with st.spinner("🔍 Performing detailed comparative analysis..."):
                        try:
                            comparison = asyncio.run(detailed_comparative_analysis(df_for_comparison))
                            st.info(f"**Comparative Analysis:**\n\n{comparison}")
                        except Exception as e:
                            st.error(f"❌ Error during comparative analysis: {e}")

        # Sub-tab 3: Agent-Based Analysis
        with comprehensive_tabs[2]:
            st.subheader("🤖 Agent-Based Analysis")
            if st.button("Run Comprehensive Agent Analysis"):
                df_for_agent = st.session_state.get("enriched_df", pd.DataFrame())
                if df_for_agent.empty:
                    st.warning("⚠️ No data to analyze. Please enrich data first.")
                else:
                    with st.spinner("🤖 Running comprehensive agent analysis..."):
                        try:
                            agent_insights = asyncio.run(run_agent_analysis(df_for_agent))
                            st.info(f"**Comprehensive Agent Analysis:**\n\n{agent_insights}")
                        except Exception as e:
                            st.error(f"❌ Error during comprehensive agent analysis: {e}")






def download_analysis_report(df: pd.DataFrame):
    """
    Generates and allows downloading of a comprehensive analysis report in Word format.
    """
    if df.empty:
        st.warning("No data available to generate a report.")
        return

    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading('Marktplaats Scraper Analysis Report', 0)

    # Add data summary
    doc.add_heading('Data Summary', 1)
    doc.add_paragraph(f"Number of listings: {len(df)}")

    # Add visualizations
    if 'Brand' in df.columns and 'CleanedPrice' in df.columns:
        doc.add_heading('Brand Prices', 1)
        brand_prices = df.groupby('Brand')['CleanedPrice'].mean().reset_index()
        brand_prices = brand_prices.sort_values('CleanedPrice', ascending=False)
        data = [['Brand', 'Average Price']] + brand_prices.values.tolist()
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(data):
            if i == 0:
                # Header row
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)
            else:
                table.add_row()
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)

        fig, ax = plt.subplots(figsize=(10, 6))
        sns.barplot(x='Brand', y='CleanedPrice', data=brand_prices, ax=ax)
        ax.set_title('Average Brand Prices')
        ax.set_xlabel('Brand')
        ax.set_ylabel('Average Price (€)')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'Sentiment' in df.columns:
        doc.add_heading('Sentiment Distribution', 1)
        sentiment_counts = df['Sentiment'].value_counts().reset_index()
        sentiment_counts.columns = ['Sentiment', 'Count']
        data = sentiment_counts.values.tolist()
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(data):
            if i == 0:
                # Header row
                table.cell(i, 0).text = "Sentiment"
                table.cell(i, 1).text = "Count"
            else:
                table.add_row()
                table.cell(i, 0).text = str(row[0])
                table.cell(i, 1).text = str(row[1])

        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(sentiment_counts['Count'], labels=sentiment_counts['Sentiment'], autopct='%1.1f%%')
        ax.axis('equal')
        ax.set_title('Sentiment Distribution')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'CleanedPrice' in df.columns:
        doc.add_heading('Price Distribution', 1)
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.histplot(data=df, x='CleanedPrice', ax=ax)
        ax.set_title('Price Distribution')
        ax.set_xlabel('Price (€)')
        ax.set_ylabel('Number of Listings')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'Usage' in df.columns and 'CleanedPrice' in df.columns:
        doc.add_heading('Usage vs. Average Price', 1)
        usage_price = df.groupby('Usage')['CleanedPrice'].mean().reset_index()
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.scatterplot(x='Usage', y='CleanedPrice', data=usage_price, ax=ax)
        ax.set_title('Usage vs. Average Price')
        ax.set_xlabel('Usage')
        ax.set_ylabel('Average Price (€)')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    # Save the document
    file_name = f"analysis_report.docx"
    doc.save(file_name)

    # Download the document
    with open(file_name, "rb") as file:
        btn = st.download_button(
            label="Download Analysis Report",
            data=file.read(),
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main() 
