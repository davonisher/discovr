######################
#  IMPORTS & SETUP   #
######################

#to do 
# 1. Add the ability to send an email alert with the docx
# 2. add the option to add specs in analysis
#3. improve the shadcn/ui
#4. add a function to send a whatsapp message with the insights


import datetime
from datetime import timedelta
import datetime
from typing import Optional
from itertools import tee
import re
import streamlit as st
import pandas as pd
import os
import sys
import asyncio
import random
import urllib.parse
import time
import json
import logging
import requests
import altair as alt
import time
import sib_api_v3_sdk
from sib_api_v3_sdk.rest import ApiException
from pprint import pprint
from playwright.async_api import async_playwright, TimeoutError
from playwright_stealth import stealth_async
import matplotlib.pyplot as plt
import seaborn as sns
import io
from docx import Document
from docx.shared import Inches
import time
import openai
from openai import OpenAI
from price_parser import Price
import streamlit as st
import pandas as pd
import time
import os
import io
import asyncio
import requests
import plotly.express as px
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# KEYS

#groq os
# KEYS
openai_api_key = os.getenv("OPENAI_API_KEY")
brevo_api_key = os.getenv("BREVO_API_KEY")
groq_api_key = os.getenv("GROQ_API_KEY")
deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")



###########################
#  FILE & PATH CONSTANTS  #
###########################

# Define the directory to store saved CSV files
SAVED_RUNS_DIR = "/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/Professioneel/Coding projects/marktplaats/data/class_data"
# Create the directory if it doesn't exist
os.makedirs(SAVED_RUNS_DIR, exist_ok=True)

# Make modules accessible if they're in a parent folder
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.join(current_dir, "..")
if project_root not in sys.path:
    sys.path.append(project_root)

#FILES
USER_AGENTS_PATH = (
    "/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/"
    "Professioneel/Coding projects/marktplaats/user-agents.txt"
)
CAT_FILE_PATH = (
    "/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/"
    "Professioneel/Coding projects/marktplaats/cat.txt"
)
CSV_EXPORT_PATH = (
    "/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/"
    "Professioneel/Coding projects/marktplaats/data/class_data/{search_query_encoded}.csv"
)
from modules.query_assistant_openai import show_query_assistant
from modules import image_search as image_search
#########################
#  SCRAPING FUNCTIONS   #
#########################

async def block_aggressively(route):
    """
    Prevent the browser from downloading images, CSS, etc.
    This speeds up scraping if you only need the HTML.
    """
    if route.request.resource_type != "document":
        await route.abort()
    else:
        await route.continue_()

def load_user_agents() -> list:
    """
    Load user agents from a text file, one per line.
    """
    with open(USER_AGENTS_PATH, "r") as f:
        return [line.strip() for line in f if line.strip()]

async def scrape_single_page(
    page_num: int,
    base_url: str,
    category: str,
    search_query_encoded: str,
    user_agent: str,
    listings: list
):
    """
    Scrape a single page of Marktplaats asynchronously.
    Gathers title, price, usage, description, link, date into 'listings' list.
    """
    url = f"{base_url}/l/{category}/q/{search_query_encoded}/p/{page_num}/"
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context(user_agent=user_agent)
        page = await context.new_page()

        # Stealth for more human-like scraping
        await stealth_async(page)
        await page.route("**/*", block_aggressively)

        try:
            await page.goto(url, timeout=60000)
            await asyncio.sleep(2)  # Wait a bit for the page to load
            await page.wait_for_load_state('networkidle')
        except TimeoutError:
            print(f"Timeout at page {page_num}, url={url}")
            await browser.close()
            return
        except Exception as e:
            print(f"Error loading page {page_num}: {e}")
            await browser.close()
            return

        # Grab listing elements
        listing_elements = await page.query_selector_all('li[class*="Listing"]')
        for element in listing_elements:
            try:
                title_el = await element.query_selector('h3[class*="Listing-title"]')
                price_el = await element.query_selector('p[class*="Listing-price"]')
                usage_el = await element.query_selector('span[class*="hz-Attribute hz-Attribute--default"]')
                desc_el = await element.query_selector('p[class*="Listing-description"]')
                link_el = await element.query_selector('a[class*="Link"]')
                date_el = await element.query_selector('span[class*="Listing-date"]')

                title_text = (await title_el.inner_text()).strip() if title_el else 'No title'
                price_text = (await price_el.inner_text()).strip() if price_el else 'No price'
                usage_text = (await usage_el.inner_text()).strip() if usage_el else 'No usage'
                desc_text = (await desc_el.inner_text()).strip() if desc_el else 'No description'
                href = (await link_el.get_attribute('href')) if link_el else ''
                full_link = (
                    "https://www.marktplaats.nl" + href
                    if href.startswith('/v/') else 'No link'
                )
                date_text = (await date_el.inner_text()).strip() if date_el else 'No date'

                listings.append({
                    'Title': title_text,
                    'Price': price_text,
                    'Description': desc_text,
                    'Link': full_link,
                    'Date': date_text,
                    'Usage': usage_text,
                })
            except Exception as ex:
                print(f"Error parsing listing on page {page_num}: {ex}")
        await browser.close()

async def scrape_marktplaats_playwright(search_query: str, category: str, max_pages: int) -> pd.DataFrame:
    """
    Main async scraping function.
    Gathers data from up to 'max_pages' pages on Marktplaats.
    """
    listings = []
    base_url = "https://www.marktplaats.nl"
    search_query_encoded = urllib.parse.quote(search_query)

    # Load a list of user agents for rotating
    user_agents = load_user_agents()

    for page_num in range(1, max_pages + 1):
        user_agent = random.choice(user_agents)
        await scrape_single_page(
            page_num=page_num,
            base_url=base_url,
            category=category,
            search_query_encoded=search_query_encoded,
            user_agent=user_agent,
            listings=listings
        )

    df = pd.DataFrame(listings)
    return df

#############################
#  CLASSIFICATION FUNCTION  #
#############################
#client = OpenAI(   base_url="https://api.deepseek.com", api_key=os.getenv("DEEPSEEK_API_KEY"),)

## TO DO: 
# improve the classification function of Deepseek API. 
# Make more requests with semaphore


async def classify_listing_with_llm2(title: str, description: str, date_str: str) -> dict:
    """
    Voorbeeld van hoe je de DeepSeek API kunt aanroepen (met 'deepseek-chat' model)
    om gestructureerde JSON terug te krijgen voor brand, model en date.
    """

    # Bouw de prompt
    user_prompt = f"""
    You are a product classifier. You receive a Title, Description, and Date from a marketplace listing.
    You must extract:
    - brand (the manufacturer or brand name)
    - model (the specific model name or identifier)
    - date (the date from the listing)

    Return the result *strictly* as JSON with keys "brand", "model", and "date".
    If you are uncertain, use "Unknown" for brand or model.

    Title: "{title}"
    Description: "{description}"
    Date: "{date_str}"

    JSON output:
    """

    # Bouw de messages (system + user), net als je voorbeeld
    system_prompt = (
        "You are an AI model tasked with extracting product information from a listing. "
        "The extracted info must be brand, model, and date, in JSON."
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    # Voorbeeld: we maken 2 of 3 retries (optioneel)
    max_retries = 3
    retry_delay = 1

    for attempt in range(max_retries):
        try:
            # Let op: DeepSeek call (NIET requests.post)!
            # Gebruik de client.chat.completions.create(...), en geef 'response_format={"type":"json_object"}' mee
            response = client.chat.completions.create(
                model="deepseek-chat",  # of een andere modelnaam als je dat hebt
                messages=messages,
                response_format={"type": "json_object"}  # Zorgt dat de output JSON is
            )

            # De API geeft iets als: response.choices[0].message.content
            if response and response.choices:
                content_str = response.choices[0].message.content
                # Parse de JSON
                result_dict = json.loads(content_str)

                final_dict = {
                    "brand": result_dict.get("brand", "Unknown"),
                    "model": result_dict.get("model", "Unknown"),
                    "date":  result_dict.get("date", date_str),
                }
                return final_dict
            else:
                logging.warning("Geen geldige respons ontvangen van DeepSeek.")
                return None

        except Exception as e:
            logging.error(f"Exception in DeepSeek call (attempt {attempt+1}): {e}")
            if attempt < max_retries - 1:
                # Wacht even en probeer opnieuw
                await asyncio.sleep(retry_delay)
                retry_delay *= 2
            else:
                return None

    return None

async def classify_listing_with_llm(title: str, description: str, date_str: str) -> dict:
    """
    Example of an async classification function that calls GPT-4 
    or any other large language model to get brand, model, and date.
    """
    prompt = f"""
    You are a product classifier. You receive a Title, Description, and Date from a marketplace listing.
    You must extract:
    - brand (the manufacturer or brand name)
    - model (the specific model name or identifier)
    - date (the date from the listing)

    Return the result *strictly* as JSON with keys "brand", "model", and "date".
    If you are uncertain, use "Unknown" for brand or model.

    Title: "{title}"
    Description: "{description}"
    Date: "{date_str}"

    JSON output:
    """

    json_data = {
        "model": "gpt-4o",  # or another model
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an AI model tasked with extracting product information from a listing. "
                    "The extracted info must be brand, model, and date, in JSON."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "functions": [
            {
                "name": "extract_product_info",
                "description": "Extracts product information from a marketplace listing.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "brand": {"type": "string"},
                        "model": {"type": "string"},
                        "date": {"type": "string"}
                    },
                    "required": ["brand", "model", "date"]
                }
            }
        ],
        "function_call": {"name": "extract_product_info"},
    }

    headers = {
        "Authorization": f"Bearer {openai.api_key}",
        "Content-Type": "application/json"
    }

    max_retries = 3
    retry_delay = 1

    for _ in range(max_retries):
        try:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=json_data,
                timeout=60
            )
            if response.status_code == 200:
                response_data = response.json()
                arguments = response_data['choices'][0]['message']['function_call']['arguments']
                result_dict = json.loads(arguments)

                final_dict = {
                    "brand": result_dict.get("brand", "Unknown"),
                    "model": result_dict.get("model", "Unknown"),
                    "date": result_dict.get("date", date_str),
                }
                return final_dict

            elif response.status_code == 429:
                logging.warning(f"Rate limit hit. Retrying in {retry_delay} seconds...")
                await asyncio.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                logging.warning(f"Got HTTP {response.status_code}. Body: {response.text}")
                return None

        except Exception as e:
            logging.error(f"Exception while calling LLM: {e}")
            await asyncio.sleep(retry_delay)
            retry_delay *= 2

    return None

def enrich_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """
    Enriches the DataFrame by classifying each listing using the LLM.
    
    Parameters:
    - df (pd.DataFrame): The raw scraped DataFrame.
    
    Returns:
    - pd.DataFrame: The enriched DataFrame with additional classification columns.
    """
    # Define the classification function
    async def classify_row_async(row):
        title = row.get("Title", "")
        description = row.get("Description", "")
        date_str = row.get("Date", "")
        usage_str = row.get("Usage", "")
        
        # Call the asynchronous classification function
        llm_result = await classify_listing_with_llm2(title, description, date_str)
        if llm_result is None:
            # If classification fails, keep it as unknown
            return {
                "Brand": "Unknown",
                "Model": "Unknown",
                "ClassifiedDate": date_str
            }
        else:
            return {
                "Brand": llm_result.get("brand", "Unknown"),
                "Model": llm_result.get("model", "Unknown"),
                "ClassifiedDate": llm_result.get("date", date_str)
            }
    
    # Run the enrichment asynchronously
    async def enrich_async():
        enriched_data = await asyncio.gather(*[classify_row_async(row) for _, row in df.iterrows()])
        return pd.DataFrame(enriched_data)
    
    # Execute the async enrichment
    enriched_df = asyncio.run(enrich_async())
    
    # Concatenate the classification results to the original DataFrame
    enriched_df = pd.concat([df.reset_index(drop=True), enriched_df.reset_index(drop=True)], axis=1)
    
    return enriched_df

############################
#  DATA HELPER FUNCTIONS   #
############################

def sanitize_filename_component(component: str) -> str:
    """
    Removes or replaces characters that are invalid in filenames.
    """
    return "".join(c for c in component if c.isalnum() or c in (' ', '_')).rstrip()

def group_prices_by_usage(df: pd.DataFrame) -> pd.DataFrame:
    if 'Usage' not in df.columns or 'CleanedPrice' not in df.columns:
        return pd.DataFrame()
    grouped = df.groupby('Usage', dropna=False)['CleanedPrice'].mean().reset_index()
    grouped.rename(columns={'CleanedPrice': 'AvgPrice'}, inplace=True)
    return grouped

#########################
#   EMAIL ALERT (demo)
#########################
###OLD
def send_email_alert(df: pd.DataFrame, insights: str):
    """
    Send market insights via email using Brevo API.
    """
    
    # Instantiate the client
    sib_api_v3_sdk.configuration.api_key['api-key'] = os.getenv("BREVO_API_KEY")
    api_instance = sib_api_v3_sdk.EmailCampaignsApi()

    # Define the campaign settings
    email_campaigns = sib_api_v3_sdk.CreateEmailCampaign(
        name="Market Insights Campaign",
        subject="Market Insights from Marktplaats Scraper",
        sender={"name": "Marktplaats Scraper", "email": "davidkakaniss@gmail.com"},
        type="classic",
        # Content that will be sent
        html_content=f"Congratulations! Here are the latest market insights:\n\n{insights}",
        # Select the recipients
        recipients={"listIds": [2, 7]},
        # Schedule the sending immediately
        scheduled_at=time.strftime("%Y-%m-%d %H:%M:%S")
    )

    # Make the call to the client
    try:
        api_response = api_instance.create_email_campaign(email_campaigns)
        pprint(api_response)
        st.success("Email campaign sent successfully!")
    except ApiException as e:
        st.error(f"Exception when calling EmailCampaignsApi->create_email_campaign: {e}")
##GOOD
## TO DO:
# 1. Add the ability to send an email alert with the docx
# 2. HTML must be right
# 3. Add the option to add files
def generate_email_html(insights: str) -> str:
    """
    Generates a styled HTML content for the email using the insights text.
    """
    html = f"""
    <html>
    <head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333333;
        }}
        h2 {{
            color: #4CAF50;
        }}
        p {{
            margin-bottom: 10px;
        }}
        ul {{
            margin-left: 20px;
        }}
    </style>
    </head>
    <body>
    <h2>Market Insights from Marktplaats Scraper</h2>
    <p>Congratulations! Here are the latest market insights:</p>
    {insights}
    <p>Best regards,<br>Marktplaats Scraper Team</p>
    </body>
    </html>
    """
    return html
def send_email_insights(insights: str, recipient_email: str, recipient_name: str):
    """
    Send market insights via a transactional email using Brevo's API.
    
    Parameters:
    - insights (str): The insights text to send.
    - recipient_email (str): Recipient's email address.
    - recipient_name (str): Recipient's name.
    """
    # Check if insights are empty
    if not insights.strip():
        st.warning("No insights available to send in the email.")
        return

    # Retrieve API key from secrets
    brevo_api_key = os.getenv("BREVO_API_KEY")
    # Define the API endpoint
    url = "https://api.brevo.com/v3/smtp/email"

    # Define the sender
    sender = {
        "name": "Marktplaats Scraper",
        "email": "davidkakaniss@gmail.com"  # Ensure this email is verified in Brevo
    }

    # Define the email content
    html_content = generate_email_html(insights)

    # Define the payload
    payload = {
        "sender": sender,
        "to": [
            {
                "email": recipient_email,
                "name": recipient_name
            }
        ],
        "subject": "Market Insights from Marktplaats Scraper",
        "htmlContent": html_content
    }

    # Define the headers
    headers = {
        "accept": "application/json",
        "api-key": brevo_api_key,
        "content-type": "application/json"
    }

    # Send the POST request
    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()  # Raise an exception for HTTP errors
        st.success("Email sent successfully!")
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error occurred: {http_err} - {response.text}")
    except Exception as err:
        st.error(f"An error occurred: {err}")

###OLD
def send_email_alert2(df: pd.DataFrame):
    """
    Send market insights via a transactional email using Brevo's API.
    """
    # Check if DataFrame is empty
    if df.empty:
        st.warning("No data available to send in the email.")
        return

    # Retrieve API key from secrets
    brevo_api_key = os.getenv("BREVO_API_KEY")

    # Define the API endpoint
    url = "https://api.brevo.com/v3/smtp/email"

    # Define the sender (you can make this dynamic as well)
    sender = {
        "name": "Marktplaats Scraper",
        "email": "davidkakaniss@gmail.com"
    }

    # Let user input recipient details
    st.subheader("Email Recipient Details")
    recipient_email = st.text_input("Recipient Email", value="david.kakanis@hotmail.com")
    recipient_name = st.text_input("Recipient Name", value="John Doe")

    if not recipient_email:
        st.warning("Please enter the recipient's email address.")
        return

    # Define the email content
    html_content = generate_email_html(df)

    # Define the payload
    payload = {
        "sender": sender,
        "to": [
            {
                "email": recipient_email,
                "name": recipient_name
            }
        ],
        "subject": "Market Insights from Marktplaats Scraper",
        "htmlContent": html_content
    }

    # Define the headers
    headers = {
        "accept": "application/json",
        "api-key": brevo_api_key,
        "content-type": "application/json"
    }

    # Send the POST request
    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()  # Raise an exception for HTTP errors
        st.success("Email sent successfully!")
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error occurred: {http_err} - {response.text}")
    except Exception as err:
        st.error(f"An error occurred: {err}")

#########################
# ANALYSIS FUNCTIONS
#########################

# 1. Sentiment Analysis
async def sentiment_analysis_agent(df: pd.DataFrame) -> pd.DataFrame:
    """
    Analyzes the sentiment of each product description and assigns a sentiment score.
    """
    if df.empty:
        st.warning("No data available for sentiment analysis.")
        return df

    if 'Description' not in df.columns:
        st.warning("No Description column found.")
        return df

    df['Sentiment'] = 'Neutral'
    df['SentimentScore'] = 0.0

    for idx, row in df.iterrows():
        description = row['Description']
        prompt = f"""
        Analyze the sentiment of the following product description and assign a sentiment score between -1 (very negative) and 1 (very positive).

        Description: "{description}"

        Provide the output strictly as JSON with keys "sentiment" and "score".
        """

        json_data = {
            "model": "gpt-4o",
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "You are an AI assistant proficient in sentiment analysis."
                    )
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
        }

        headers = {
            "Authorization": f"Bearer {openai.api_key}",
            "Content-Type": "application/json"
        }

        max_retries = 3
        retry_delay = 1

        for _ in range(max_retries):
            try:
                response = requests.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers=headers,
                    json=json_data,
                    timeout=60
                )
                if response.status_code == 200:
                    response_data = response.json()
                    sentiment_json = response_data['choices'][0]['message']['content'].strip()
                    sentiment_dict = json.loads(sentiment_json)

                    df.at[idx, 'Sentiment'] = sentiment_dict.get("sentiment", "Neutral")
                    df.at[idx, 'SentimentScore'] = float(sentiment_dict.get("score", 0.0))
                    break
                elif response.status_code == 429:
                    logging.warning(f"Rate limit hit. Retrying in {retry_delay} seconds...")
                    await asyncio.sleep(retry_delay)
                    retry_delay *= 2
                    continue
                else:
                    logging.warning(f"Unexpected HTTP status: {response.status_code}. Body: {response.text}")
                    df.at[idx, 'Sentiment'] = "Error"
                    df.at[idx, 'SentimentScore'] = 0.0
                    break
            except Exception as e:
                logging.error(f"Exception during sentiment analysis: {e}")
                await asyncio.sleep(retry_delay)
                retry_delay *= 2
        else:
            df.at[idx, 'Sentiment'] = "Error"
            df.at[idx, 'SentimentScore'] = 0.0

    return df

# 2. Best Deal Recommendation
async def best_deal_recommendation(df: pd.DataFrame) -> str:
    """
    Identifies the listing that offers the best value for money considering price, brand reputation,
    usage condition, and model popularity.
    """
    if df.empty:
        return "No data available for best deal recommendation."

    # Ensure CleanedPrice is present
    if 'CleanedPrice' not in df.columns:
        df['CleanedPrice'] = df['Price'].apply(clean_price)

    # Prepare data summary for the LLM
    top_brands = df['Brand'].value_counts().head(5).to_dict()
    top_models = df['Model'].value_counts().head(5).to_dict()
    usage_conditions = df['Usage'].value_counts().to_dict()

    summary = f"""
    You are an intelligent agent assisting in identifying the best deal from marketplace listings.

    **Data Summary:**
    - Total Listings: {len(df)}
    - Average Price: €{df['CleanedPrice'].mean():.2f}
    - Median Price: €{df['CleanedPrice'].median():.2f}
    - Top Brands: {', '.join([f"{brand} ({count})" for brand, count in top_brands.items()])}
    - Top Models: {', '.join([f"{model} ({count})" for model, count in top_models.items()])}
    - Usage Conditions: {', '.join([f"{usage} ({count})" for usage, count in usage_conditions.items()])}

    **Criteria for Best Deal:**
    - Price
    - Brand Reputation
    - Usage Condition
    - Model Popularity

    **Please identify and describe the best deal listing based on the above criteria. Provide the Title, Price, Brand, Model, Usage, and Link.**
    """

    json_data = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an AI assistant skilled in data analysis and identifying optimal deals."
                )
            },
            {
                "role": "user",
                "content": summary
            }
        ],
        "temperature": 0.5,
    }

    headers = {
        "Authorization": f"Bearer {openai.api_key}",
        "Content-Type": "application/json"
    }

    max_retries = 3
    retry_delay = 1

    for _ in range(max_retries):
        try:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=json_data,
                timeout=60
            )
            if response.status_code == 200:
                response_data = response.json()
                recommendation = response_data['choices'][0]['message']['content'].strip()
                return recommendation
            elif response.status_code == 429:
                logging.warning(f"Rate limit hit. Retrying in {retry_delay} seconds...")
                await asyncio.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                logging.warning(f"Unexpected HTTP status: {response.status_code}. Body: {response.text}")
                return "An error occurred while identifying the best deal."
        except Exception as e:
            logging.error(f"Exception during best deal recommendation: {e}")
            await asyncio.sleep(retry_delay)
            retry_delay *= 2

    return "Failed to identify the best deal after multiple attempts."

# 3. Price Trend Analysis
async def price_trend_analysis(df: pd.DataFrame) -> str:
    """
    Analyzes pricing trends for different brands and models over time.
    """
    if df.empty:
        return "No data available for price trend analysis."

    # Ensure CleanedPrice and Date are present
    if 'CleanedPrice' not in df.columns:
        df['CleanedPrice'] = df['Price'].apply(clean_price)
    if 'Date' not in df.columns:
        st.warning("No Date column found.")
        return "No date information available for trend analysis."

    # Convert Date to datetime
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce')

    # Sort by Date
    df = df.sort_values('Date')

    # Analyze trends for top 5 brands
    top_brands = df['Brand'].value_counts().head(5).index.tolist()

    trend_summary = ""

    for brand in top_brands:
        brand_df = df[df['Brand'] == brand]
        if brand_df['Date'].isnull().all():
            trend = "No date information available."
        else:
            brand_df = brand_df.dropna(subset=['Date'])
            brand_df = brand_df.set_index('Date').resample('M')['CleanedPrice'].mean().reset_index()
            price_diff = brand_df['CleanedPrice'].iloc[-1] - brand_df['CleanedPrice'].iloc[0]
            if price_diff > 0:
                trend = "Increasing"
            elif price_diff < 0:
                trend = "Decreasing"
            else:
                trend = "Stable"
        trend_summary += f"**{brand}**: {trend}\n\n"

    return trend_summary

# 4. Market Insights
async def market_insights_analysis(df: pd.DataFrame) -> str:
    """
    Provides comprehensive market insights based on the scraped data.
    """
    if df.empty:
        return "No data available for market insights analysis."

    # Prepare summary
    total_listings = len(df)
    avg_price = df['CleanedPrice'].mean() if 'CleanedPrice' in df.columns else 0.0
    median_price = df['CleanedPrice'].median() if 'CleanedPrice' in df.columns else 0.0
    top_brands = df['Brand'].value_counts().head(5).to_dict()
    top_models = df['Model'].value_counts().head(5).to_dict()
    usage_conditions = df['Usage'].value_counts().to_dict()

    summary = f"""
    You are an AI assistant specialized in market analysis. Based on the provided data, please provide comprehensive market insights.

    **Data Summary:**
    - Total Listings: {total_listings}
    - Average Price: €{avg_price:.2f}
    - Median Price: €{median_price:.2f}
    - Top Brands: {', '.join([f"{brand} ({count})" for brand, count in top_brands.items()])}
    - Top Models: {', '.join([f"{model} ({count})" for model, count in top_models.items()])}
    - Usage Conditions: {', '.join([f"{usage} ({count})" for usage, count in usage_conditions.items()])}

    **Please provide your market insights below:**
    """

    json_data = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an AI assistant specialized in market analysis and providing comprehensive insights."
                )
            },
            {
                "role": "user",
                "content": summary
            }
        ],
        "temperature": 0.7,
    }

    headers = {
        "Authorization": f"Bearer {openai.api_key}",
        "Content-Type": "application/json"
    }

    max_retries = 3
    retry_delay = 1

    for _ in range(max_retries):
        try:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=json_data,
                timeout=60
            )
            if response.status_code == 200:
                response_data = response.json()
                insights = response_data['choices'][0]['message']['content'].strip()
                return insights
            elif response.status_code == 429:
                logging.warning(f"Rate limit hit. Retrying in {retry_delay} seconds...")
                await asyncio.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                logging.warning(f"Unexpected HTTP status: {response.status_code}. Body: {response.text}")
                return "An error occurred while generating market insights."
        except Exception as e:
            logging.error(f"Exception during market insights analysis: {e}")
            await asyncio.sleep(retry_delay)
            retry_delay *= 2

    return "Failed to generate market insights after multiple attempts."

# 5. Market Demand Forecasting
async def market_demand_forecasting(df: pd.DataFrame) -> str:
    """
    Predicts future demand or price trends for specific categories or brands based on historical data.
    """
    if df.empty:
        return "No data available for market demand forecasting."

    # Ensure CleanedPrice and Date are present
    if 'CleanedPrice' not in df.columns:
        df['CleanedPrice'] = df['Price'].apply(clean_price)
    if 'Date' not in df.columns:
        st.warning("No Date column found.")
        return "Date information is required for forecasting."

    # Convert Date to datetime
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce')

    # Sort by Date
    df = df.sort_values('Date')

    # Analyze for top 5 brands
    top_brands = df['Brand'].value_counts().head(5).index.tolist()

    forecasting_summary = ""

    for brand in top_brands:
        brand_df = df[df['Brand'] == brand]
        if brand_df['Date'].isnull().all():
            forecast = "No date information available."
        else:
            brand_df = brand_df.dropna(subset=['Date'])
            brand_df = brand_df.set_index('Date').resample('M')['CleanedPrice'].mean().reset_index()
            # Prepare data for LLM
            historical_data = brand_df.to_dict(orient='records')
            prompt = f"""
            You are an expert in market demand forecasting. Based on the historical monthly average prices below for the brand "{brand}", predict the average prices for the next 3 months.

            Historical Data:
            {historical_data}

            **Please provide your forecast below:**
            """

            json_data = {
                "model": "gpt-4",
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "You are an AI assistant specialized in market demand forecasting and price trend analysis."
                        )
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.6,
            }

            headers = {
                "Authorization": f"Bearer {openai.api_key}",
                "Content-Type": "application/json"
            }

            max_retries = 3
            retry_delay = 1

            for _ in range(max_retries):
                try:
                    response = requests.post(
                        "https://api.openai.com/v1/chat/completions",
                        headers=headers,
                        json=json_data,
                        timeout=60
                    )
                    if response.status_code == 200:
                        response_data = response.json()
                        forecast = response_data['choices'][0]['message']['content'].strip()
                        forecasting_summary += f"**{brand}** Forecast:\n{forecast}\n\n"
                        break
                    elif response.status_code == 429:
                        logging.warning(f"Rate limit hit. Retrying in {retry_delay} seconds...")
                        await asyncio.sleep(retry_delay)
                        retry_delay *= 2
                        continue
                    else:
                        logging.warning(f"Unexpected HTTP status: {response.status_code}. Body: {response.text}")
                        forecasting_summary += f"**{brand}**: An error occurred during forecasting.\n\n"
                        break
                except Exception as e:
                    logging.error(f"Exception during market demand forecasting: {e}")
                    await asyncio.sleep(retry_delay)
                    retry_delay *= 2
            else:
                forecasting_summary += f"**{brand}**: Failed to forecast after multiple attempts.\n\n"

    return forecasting_summary

############################
#  DATA HELPER FUNCTIONS   #
############################

def is_valid_email(email: str) -> bool:
    """
    Validates the email format.
    """
    regex = r'^[\w\.-]+@[\w\.-]+\.\w+$'
    return re.match(regex, email) is not None
def get_class_data_files():
    """
    Retrieves a list of files from the class_data directory.
    
    Returns:
    - List of filenames if the directory exists.
    - Empty list if the directory does not exist.
    """
    folder_path = "/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/Professioneel/Coding projects/marktplaats/data/class_data"
    try:
        files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
        return files
    except FileNotFoundError:
        st.error("❌ class_data folder not found. Please check the file path.")
        return []

# TO DO: 
# FIX FUCNITON
# Implement function to get raw data files

def get_raw_data_files():
    """
    Retrieves a list of files from the class_data directory.

    Returns:
    - List of filenames if the directory exists.
    - Empty list if the directory does not exist.
    """
    folder_path = "/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/Professioneel/Coding projects/marktplaats/data/raw_data"
    try:
        files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
        return files
    except FileNotFoundError:
        st.error("❌ class_data folder not found. Please check the file path.")
        return []
def get_category_options():
    """
    Fetches category options from the cat.txt file.
    Returns a list of categories.
    """
    try:
        with open(CAT_FILE_PATH, "r") as f:
            category_options = [line.strip() for line in f if line.strip()]
        return category_options
    except FileNotFoundError:
        st.error("❌ cat.txt not found. Check the file path.")
        return []

#########################
# DATA PREPROCESSING FUNCTIONS
#########################

def clean_price(price_str: str) -> float:
    """
    Extracts the numeric amount from a price string using the price-parser library.

    Parameters:
    - price_str (str): Raw price string (e.g., "22,90 €", "$119.00").

    Returns:
    - float: Cleaned price amount. Returns 0.0 if the price can't be parsed.
    """
    if not isinstance(price_str, str):
        return 0.0

    # Use price-parser to parse the price
    price = Price.fromstring(price_str)

    # Return the amount as a float, defaulting to 0.0 if parsing fails
    return price.amount_float or 0.0



def parse_date(date_str: str) -> Optional[datetime.date]:
    date_str = str(date_str).strip().lower()
    today = datetime.date.today()  # Gebruik datetime.date.today()

    # Handle relative dates
    relative_dates = {
        'vandaag': today,
        'gisteren': today - datetime.timedelta(days=1),
        'eergisteren': today - datetime.timedelta(days=2),
    }
    if date_str in relative_dates:
        return relative_dates[date_str]

    # Handle absolute dates with Dutch month abbreviations
    month_mapping = {
        'jan': 1, 'feb': 2, 'mrt': 3, 'apr': 4, 'mei': 5, 'jun': 6,
        'jul': 7, 'aug': 8, 'sep': 9, 'okt': 10, 'nov': 11, 'dec': 12
    }
    match = re.match(r'(\d{1,2})\s+([a-z]{3})\s+(\d{2})', date_str)
    if match:
        day, month_abbr, year_suffix = match.groups()
        month = month_mapping.get(month_abbr)
        if month:
            year = 2000 + int(year_suffix)  # Assuming 2000+ for '25'
            try:
                return datetime.date(year, month, int(day))
            except ValueError:
                pass
    return None

def preprocess_enriched_df(df: pd.DataFrame) -> pd.DataFrame:
    """
    Preprocesses the enriched DataFrame by converting date columns to datetime.
    
    Parameters:
    - df (pd.DataFrame): The enriched DataFrame.
    
    Returns:
    - pd.DataFrame: The preprocessed DataFrame.
    """
    if 'Date' in df.columns:
        df['Date'] = df['Date'].apply(parse_date)
    
    if 'Price' in df.columns:
        df['CleanedPrice'] = df['Price'].apply(clean_price)
    return df


#########################
# DASHBOARD DISPLAY
#########################

def show_dashboard(df: pd.DataFrame):
    """
    Display multiple charts and stats:
      1) Pie chart of average price by usage.
      2) Pie chart of the most popular brand.
      3) Pie chart of the most popular model.
      4) Bar chart of brand prices.
      5) Price Distribution Histogram.
      6) Sentiment Distribution.
    """
    st.header("Dashboard")

    # 1) Average Price by Usage (Pie Chart)
    st.subheader("Average Price by Usage")
    if 'CleanedPrice' in df.columns and 'Usage' in df.columns:
        usage_groups = df.groupby('Usage', dropna=False)['CleanedPrice'].mean().reset_index()
        usage_groups.rename(columns={'CleanedPrice': 'AvgPrice'}, inplace=True)

        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(usage_groups['AvgPrice'], labels=usage_groups['Usage'], autopct='%1.1f%%')
        ax.axis('equal')
        ax.set_title('Average Price by Usage')
        st.pyplot(fig)
    else:
        st.write("No usage or price data available.")

    # 2) Most Popular Brand (Pie Chart)
    st.subheader("Most Popular Brand")
    if 'Brand' in df.columns:
        brand_counts = df['Brand'].value_counts().reset_index()
        brand_counts.columns = ['Brand', 'Count']

        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(brand_counts['Count'], labels=brand_counts['Brand'], autopct='%1.1f%%')
        ax.axis('equal')
        ax.set_title('Most Popular Brand')
        st.pyplot(fig)
    else:
        st.write("No brand data available.")

    # 3) Most Popular Model (Pie Chart)
    st.subheader("Most Popular Model")
    if 'Model' in df.columns:
        model_counts = df['Model'].value_counts().reset_index()
        model_counts.columns = ['Model', 'Count']

        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(model_counts['Count'], labels=model_counts['Model'], autopct='%1.1f%%')
        ax.axis('equal')
        ax.set_title('Most Popular Model')
        st.pyplot(fig)
    else:
        st.write("No model data available.")

    # 4) Brand Prices (Bar Chart)
    st.subheader("Average Brand Prices")
    if 'Brand' in df.columns and 'CleanedPrice' in df.columns:
        brand_prices = df.groupby('Brand')['CleanedPrice'].mean().reset_index()
        brand_prices.columns = ['Brand', 'AvgPrice']

        fig, ax = plt.subplots(figsize=(10, 6))
        sns.barplot(x='Brand', y='AvgPrice', data=brand_prices, ax=ax)
        ax.set_title('Average Brand Prices')
        ax.set_xlabel('Brand')
        ax.set_ylabel('Average Price (€)')
        st.pyplot(fig)
    else:
        st.write("No brand or price data available.")

    # 5) Price Distribution (Histogram)
    st.subheader("Price Distribution")
    if 'CleanedPrice' in df.columns:
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.histplot(data=df, x='CleanedPrice', ax=ax)
        ax.set_title('Price Distribution')
        ax.set_xlabel('Price (€)')
        ax.set_ylabel('Number of Listings')
        st.pyplot(fig)
    else:
        st.write("No price data available.")

    # 6) Sentiment Distribution (Pie Chart)
    st.subheader("Sentiment Distribution")
    if 'Sentiment' in df.columns:
        sentiment_counts = df['Sentiment'].value_counts().reset_index()
        sentiment_counts.columns = ['Sentiment', 'Count']

        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(sentiment_counts['Count'], labels=sentiment_counts['Sentiment'], autopct='%1.1f%%')
        ax.axis('equal')
        ax.set_title('Sentiment Distribution')
        st.pyplot(fig)
    else:
        st.write("No sentiment data available.")

    # 7) Additional Visualizations (Optional)
    st.subheader("Additional Visualizations")

    # Example: Usage vs. Average Price
    if 'Usage' in df.columns and 'CleanedPrice' in df.columns:
        usage_price = df.groupby('Usage')['CleanedPrice'].mean().reset_index()

        fig, ax = plt.subplots(figsize=(10, 6))
        sns.scatterplot(x='Usage', y='CleanedPrice', data=usage_price, ax=ax)
        ax.set_title('Usage vs. Average Price')
        ax.set_xlabel('Usage')
        ax.set_ylabel('Average Price (€)')
        st.pyplot(fig)
    else:
        st.write("Insufficient data for Usage vs. Average Price visualization.")

#########################
# MAIN SCRAPER ORCHESTRA
#########################

async def orchestrate_scraper(search_query: str, category: str, max_pages: int) -> pd.DataFrame:
    """
    The main async pipeline:
    1. Scrape the listings via Playwright.
    2. Enrich each listing with brand/model/date via LLM.
    3. Perform Sentiment Analysis.
    4. Return a DataFrame with the final data.
    """
    # 1. SCRAPE
    df_listings = await scrape_marktplaats_playwright(search_query, category, max_pages)
    if df_listings.empty:
        return df_listings

    # 2. ENRICH
    enriched_rows = []
    for _, row in df_listings.iterrows():
        title = row.get("Title", "")
        description = row.get("Description", "")
        date_str = row.get("Date", "")
        usage_str = row.get("Usage", "")

        llm_result = await classify_listing_with_llm(title, description, date_str)
        if llm_result is None:
            # If classification fails, keep it as unknown
            llm_result = {
                "brand": "Unknown",
                "model": "Unknown",
                "date": date_str,
            }

        row_dict = {
            "Title": title,
            "Price": row.get("Price", ""),
            "Description": description,
            "Link": row.get("Link", ""),
            "Date": date_str,
            "Usage": usage_str,
            "Brand": llm_result["brand"],
            "Model": llm_result["model"],
            "ClassifiedDate": llm_result["date"],
        }
        enriched_rows.append(row_dict)

    enriched_df = pd.DataFrame(enriched_rows)

    # 3. SENTIMENT ANALYSIS
    enriched_df = await sentiment_analysis_agent(enriched_df)

    return enriched_df
#########################
# AGENT ANALYSIS FUNCTIONS
#########################

async def run_agent_analysis(df: pd.DataFrame) -> str:
    """
    Comprehensive agent analysis combining multiple insights.
    """
    insights = ""

    # Best Deal Recommendation
    best_deal = await best_deal_recommendation(df)
    insights += f"**Best Deal Recommendation:**\n{best_deal}\n\n"

    # Price Trend Analysis
    price_trend = await price_trend_analysis(df)
    insights += f"**Price Trend Analysis:**\n{price_trend}\n\n"

    # Market Insights
    market_insights = await market_insights_analysis(df)
    insights += f"**Market Insights:**\n{market_insights}\n\n"

    # Market Demand Forecasting
    demand_forecasting = await market_demand_forecasting(df)
    insights += f"**Market Demand Forecasting:**\n{demand_forecasting}\n\n"

    return insights

#########################
# DETAILED COMPARATIVE ANALYSIS
#########################

async def detailed_comparative_analysis(df: pd.DataFrame) -> str:
    """
    Uses an LLM to compare similar listings and highlight the best options based on multiple criteria.
    """
    if df.empty:
        return "No data available for comparative analysis."

    # Ensure CleanedPrice is present
    if 'CleanedPrice' not in df.columns:
        df['CleanedPrice'] = df['Price'].apply(clean_price)

    # Select top 10 cheapest listings as examples
    sample_df = df.nsmallest(10, 'CleanedPrice')

    # Convert sample data to a readable format
    listings_info = ""
    for idx, row in sample_df.iterrows():
        listings_info += f"""
        **Listing {idx + 1}:**
        - **Title:** {row['Title']}
        - **Price:** {row['Price']}
        - **Brand:** {row['Brand']}
        - **Model:** {row['Model']}
        - **Usage:** {row['Usage']}
        - **Link:** {row['Link']}
        - **Date:** {row['Date']}
        - **Sentiment:** {row.get('Sentiment', 'Neutral')} (Score: {row.get('SentimentScore', 0.0)})
        """

    prompt = f"""
    You are an expert market analyst. Compare the following marketplace listings and recommend the best options based on price, brand reputation, model popularity, usage condition, and sentiment score.

    {listings_info}

    **Provide your analysis and recommendations below:**
    """

    json_data = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an AI assistant specialized in market analysis and comparative evaluations."
                )
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.7,
    }

    headers = {
        "Authorization": f"Bearer {openai.api_key}",
        "Content-Type": "application/json"
    }

    max_retries = 3
    retry_delay = 1

    analysis = ""

    for _ in range(max_retries):
        try:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=json_data,
                timeout=60
            )
            if response.status_code == 200:
                response_data = response.json()
                analysis = response_data['choices'][0]['message']['content'].strip()
                break
            elif response.status_code == 429:
                logging.warning(f"Rate limit hit. Retrying in {retry_delay} seconds...")
                await asyncio.sleep(retry_delay)
                retry_delay *= 2
                continue
            else:
                logging.warning(f"Unexpected HTTP status: {response.status_code}. Body: {response.text}")
                analysis = "An error occurred during comparative analysis."
                break
        except Exception as e:
            logging.error(f"Exception during comparative analysis: {e}")
            await asyncio.sleep(retry_delay)
            retry_delay *= 2
    else:
        analysis = "Failed to perform comparative analysis after multiple attempts."

    return analysis


import streamlit as st
import asyncio
import uuid
from typing import List

###################################################
# Voorbeeld van de background-functie, asynchroon #
###################################################


async def run_background_analysis(job_id: str, df: pd.DataFrame, email: str, analysis_types: List[str]):
    """
    Asynchrone functie die meerdere analyses draait op df en na afloop
    een email stuurt. We halen job-info uit st.session_state["analysis_jobs"][job_id].
    """
    # Haal de job op uit de sessie (optioneel)
    job_info = st.session_state["analysis_jobs"][job_id]
    job_info["status"] = "running"
    st.session_state["analysis_jobs"][job_id] = job_info
    
    try:
        # 1. Doe (eventueel) checks of df niet leeg is
        if df.empty:
            job_info["status"] = "error"
            job_info["error_message"] = "DataFrame is empty."
            st.session_state["analysis_jobs"][job_id] = job_info
            return

        # 2. Loop over de geselecteerde analyses
        results = {}
        for analysis in analysis_types:
            if analysis == "Market Demand":
                # results["market_demand"] = ... # roep je analysis-functie aan
                await asyncio.sleep(1)  # demo: simulate time
                results["market_demand"] = "This is a demo result for Market Demand"
            elif analysis == "Comparative Analysis":
                # results["comparative"] = ...
                await asyncio.sleep(1)
                results["comparative"] = "This is a demo result for Comparative Analysis"
            elif analysis == "Agent Analysis":
                # results["agent"] = ...
                await asyncio.sleep(1)
                results["agent"] = "This is a demo result for Agent Analysis"
        
        # 3. Sla de resultaten op in job_info
        job_info["results"] = results
        job_info["status"] = "completed"
        st.session_state["analysis_jobs"][job_id] = job_info
        
        # 4. Stuur eventueel een mail
        # send_email_report(to=email, content=results)  # Zelf bouwen

    except Exception as e:
        job_info["status"] = "error"
        job_info["error_message"] = str(e)
        st.session_state["analysis_jobs"][job_id] = job_info


def queue_background_analysis():
    """
    Laat de gebruiker analyses in een achtergrondjob queue'en.
    Zodra de gebruiker op 'Queue Analysis' klikt, wordt een job_id aangemaakt
    en run_background_analysis asynchroon aangeroepen.
    
    Vereisten:
    - st.session_state["enriched_df"] moet bestaan (via 'Enrich Data' stap).
    - in streamlit: import uuid, datetime, etc.
    - run_background_analysis is async en draait analyses -> mails user.
    """
    if "enriched_df" not in st.session_state:
        st.warning("⚠️ No data available for analysis. Please enrich data first.")
        return

    st.subheader("Queue a Background Analysis")

    with st.form("background_analysis_form", clear_on_submit=False):
        email = st.text_input("Email address for notifications:", value="", key="queue_bg_email")
        analysis_types = st.multiselect(
            "Select analyses to run:",
            ["Market Demand", "Comparative Analysis", "Agent Analysis"],
            default=["Market Demand", "Comparative Analysis", "Agent Analysis"],
            key="queue_bg_analyses"
        )
        submit = st.form_submit_button("Queue Analysis")

    if submit:
        if not email:
            st.error("❌ Please provide an email address.")
            return
        
        # Eventueel extra check of df niet leeg is
        df = st.session_state["enriched_df"]
        if df.empty:
            st.warning("⚠️ There's data in 'enriched_df', but it's empty. Aborting.")
            return
        
        st.info(f"⏳ Queueing your analysis job for {email} ...")

        try:
            # 1. Maak job_id
            job_id = str(uuid.uuid4())
            
            # 2. Sla job-info op in session_state
            if "analysis_jobs" not in st.session_state:
                st.session_state["analysis_jobs"] = {}
            st.session_state["analysis_jobs"][job_id] = {
                "email": email,
                "analysis_types": analysis_types,
                "status": "queued",
                "timestamp": datetime.now().isoformat(),
                "results": {}
            }
            
            # 3. Start de achtergrondtaak (async)
            asyncio.create_task(run_background_analysis(
                job_id=job_id,
                df=df,
                email=email,
                analysis_types=analysis_types
            ))
            
            st.success(f"✅ Analysis queued successfully (job: {job_id}). "
                       f"You will receive an email at {email} when done.")
            
        except Exception as e:
            st.error(f"Error queueing analysis: {str(e)}")


def show_analysis_queue_status():
    """
    Optionele functie om de status van bestaande jobs te tonen.
    """
    if "analysis_jobs" not in st.session_state or not st.session_state["analysis_jobs"]:
        st.warning("No analysis jobs queued.")
        return
    
    st.subheader("Analysis Jobs Status")
    jobs = st.session_state["analysis_jobs"]
    for job_id, job_info in jobs.items():
        st.write(f"**Job ID:** {job_id}")
        st.write(f"- Email: {job_info['email']}")
        st.write(f"- Analyses: {job_info['analysis_types']}")
        st.write(f"- Status: {job_info['status']}")
        if job_info['status'] == 'completed':
            st.write(f"- Results: {job_info['results']}")
        elif job_info['status'] == 'error':
            st.write(f"- Error: {job_info.get('error_message', 'Unknown error')}")
        st.write("---")

import plotly.express as px


# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   ANY OTHER IMPORTS / HELPER FUNCTIONS YOU ALREADY HAD
#   (e.g. get_category_options, get_raw_data_files, etc.)
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

def sanitize_filename_component(value: str) -> str:
    """Helper to sanitize a filename component."""
    # (Your existing implementation)
    return "".join(c for c in value if c.isalnum() or c in (" ", "_", "-")).rstrip()

def load_orchestrate_and_display():
    """
    Enhanced UI with tabs for better organization and user experience.
    """
    st.set_page_config(page_title="A Trader's Edge", layout="wide", initial_sidebar_state="expanded")
    
    # Title and Description
    st.title("🔍 A Trader's Edge")
    st.markdown("""
    **Unparalleled Insights:** Data-Driven Analysis Delivered Directly to Your Inbox.

    This application provides a comprehensive approach to market analysis. By gathering data from online marketplaces through web scraping, 
    we process and analyze this information using advanced techniques.

    Our algorithms uncover valuable insights into market trends, pricing, and consumer sentiment. 
    These insights are compiled and presented in an easy-to-understand format, allowing you to make informed decisions.

    **Additionally**, we deliver these insights directly to your inbox, ensuring you never miss important market developments. 
    Stay informed, stay competitive, and let our data-driven analysis be your guide to success.
    """)    
    # Create tabs
    tabs = st.tabs([
        "Getting Data", 
        "Enrich Data",
        "Analysis",
        "Email Insights",
        "Comprehensive Analysis",
        "Dashboard",
        "Background Analysis",
        "Image Search"
    ])

    # Tab 1: Getting Data
    with tabs[0]:
        tab_1_getting_data()

    # Tab 2: Enrich Data  
    with tabs[1]:
        tab_2_enrich_data()

    # Tab 3: Analysis
    with tabs[2]:
        tab_3_analysis()

    # Tab 4: Email Insights
    with tabs[3]:
        tab_4_email_insights()

    # Tab 5: Comprehensive Analysis
    with tabs[4]:
        tab_5_comprehensive_analysis()

    # Tab 6: Dashboard
    with tabs[5]:
        tab_6_dashboard()

    # Tab 7: Background Analysis
    with tabs[6]:
        tab_7_background_analysis()

    # Tab 8: Image Search
    with tabs[7]:
        tab_8_image_search()






def sanitize_filename_component(value: str) -> str:
    """Helper to sanitize a filename component."""
    # (Your existing implementation)
    return "".join(c for c in value if c.isalnum() or c in (" ", "_", "-")).rstrip()

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   TAB 1: SCRAPE DATA
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

# -- Je bestaande imports en code hier --
# from query_assistant import show_query_assistant

def tab_1_getting_data():
    """
    Code originally found under 'with tabs[0]:'
    """
    st.header("getting data")

    # ====================
    # Optional: Deepseek assistent
    # ====================
    with st.expander("Gebruik Zoekhulp"):
        chosen_query = show_query_assistant()
        if chosen_query:
            st.session_state["assistant_query"] = chosen_query
            st.success(f"Gekozen zoekopdracht uit assistent: {chosen_query}")

    # ====================
    # Jouw bestaande formulier
    # ====================
    with st.form(key='scrape_form'):
        col1, col2 = st.columns([3, 1])
        
        with col1:
            # Als we de assistent-query willen overnemen in 'search_query':
            default_query = st.session_state.get("assistant_query", "antiek")
            search_query = st.text_input("✅ Voer je zoekterm in:", value=default_query)

            category = st.selectbox("📂 Categorie", get_category_options())
            max_pages = st.number_input("📄 Aantal pagina's (1-200)", min_value=1, max_value=200, value=2)
        
        with col2:
            st.empty()  # Placeholder for alignment

        submit_scrape = st.form_submit_button(label="Zoek")
        estimated_time = max_pages * 7
        st.info(f"⏱️ Estimated wait time: {estimated_time} seconds")

    if submit_scrape:
        with st.spinner("🕵️‍♂️ Researching..."):
            try:
                # Scrape asynchronously
                df_scraped = asyncio.run(scrape_marktplaats_playwright(search_query, category, max_pages))
            except Exception as e:
                st.error(f"❌ Error during scraping: {e}")
                return

            if df_scraped.empty:
                st.warning("⚠️ Geen listings gevonden.")
            else:
                st.success("✅ Done. Data preview:")
                st.dataframe(df_scraped.head(10))

                # Save to session_state and CSV
                st.session_state["raw_scraped_df"] = df_scraped

                # Generate a unique filename with timestamp
                timestamp = int(time.time())
                sanitized_search_query = sanitize_filename_component(search_query)
                sanitized_category = sanitize_filename_component(category)
                csv_filename = f"Raw_{sanitized_search_query}_{sanitized_category}_{max_pages}_{timestamp}.csv"
                csv_path = os.path.join(SAVED_RUNS_DIR, csv_filename)
                df_scraped.to_csv(csv_path, index=False)
                st.info(f"📁 Raw data exported to CSV at: {csv_path}")

                if "saved_data" not in st.session_state:
                    st.session_state["saved_data"] = {}
                run_label = f"Raw_{sanitized_search_query}_{sanitized_category}_{max_pages}_{timestamp}"
                st.session_state["saved_data"][run_label] = csv_path


# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   TAB 2: ENRICH DATA
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def tab_2_enrich_data():
    """
    Code originally found under 'with tabs[1]:'
    """
    st.header("🛠️ Enrich Data")

    saved_data = st.session_state.get("saved_data", {})
    raw_runs = {k: v for k, v in saved_data.items() if k.startswith("Raw_")}

    # Add raw data selection from the specified folder
    raw_data_files = get_raw_data_files()
    if raw_data_files:
        selected_raw_files = st.multiselect("✅ Select raw data files to enrich:", options=raw_data_files)
        if selected_raw_files:
            if st.button("📥 Load Selected Raw Data Files"):
                loaded_files = []
                for file in selected_raw_files:
                    file_path = os.path.join("/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/Professioneel/Coding projects/marktplaats/data/raw_data", file)
                    try:
                        df_raw = pd.read_csv(file_path)
                        # Merge with existing raw_df or create new one
                        if "raw_df" in st.session_state:
                            st.session_state["raw_df"] = pd.concat([st.session_state["raw_df"], df_raw], ignore_index=True)
                        else:
                            st.session_state["raw_df"] = df_raw
                        loaded_files.append(file)
                    except Exception as e:
                        st.error(f"❌ Error loading {file}: {e}")
                if loaded_files:
                    st.success(f"✅ Loaded raw data files: {', '.join(loaded_files)}")
                    st.dataframe(st.session_state["raw_df"].head(10))
    else:
        st.warning("⚠️ No raw data files available in the folder.")

    # Button for preprocessing
    if "raw_df" in st.session_state:
        if st.button("🔄 Preprocess Data"):
            try:
                df_raw = st.session_state["raw_df"]
                print("Raw DataFrame:")
                print(df_raw)

                # Ensure necessary columns exist
                required_columns = ["Title", "Description", "Date", "Usage", "Price"]
                if not all(col in df_raw.columns for col in required_columns):
                    st.error(f"❌ Raw data is missing required columns: {required_columns}")
                    return

                # Preprocess the DataFrame
                preprocessed_df = preprocess_enriched_df(df_raw)
                st.session_state["preprocessed_df"] = preprocessed_df
                print("Preprocessed DataFrame:")
                print(preprocessed_df)
                
                # Verify price preprocessing
                if 'CleanedPrice' in preprocessed_df.columns:
                    print("Sample of Price vs CleanedPrice:")
                    print(preprocessed_df[['Price', 'CleanedPrice']].head())
                
                if preprocessed_df.empty:
                    st.warning("⚠️ No listings found after preprocessing.")
                else:
                    st.success("✅ Preprocessing complete. Data preview:")
                    st.dataframe(preprocessed_df.head(10))
            except Exception as e:
                st.error(f"❌ Error during preprocessing: {e}")
                print(f"Preprocessing error details: {str(e)}")

        # Button for enriching
        if st.button("🛠️ Enrich Data"):
            if "preprocessed_df" not in st.session_state:
                st.error("❌ Please preprocess the data first.")
            else:
                with st.spinner("🛠️ Enriching data via LLM..."):
                    try:
                        # Enrich the DataFrame
                        enriched_df = enrich_dataframe(st.session_state["preprocessed_df"])
                        st.session_state["enriched_df"] = enriched_df

                        if enriched_df.empty:
                            st.warning("⚠️ No listings found after enrichment.")
                        else:
                            st.success("✅ Enrichment complete. Data preview:")
                            st.dataframe(enriched_df.head(10))

                            # Generate a unique filename with timestamp
                            timestamp = int(time.time())
                            sanitized_search_query = sanitize_filename_component(search_query)
                            sanitized_category = sanitize_filename_component(category)
                            csv_filename = f"Enriched_{sanitized_search_query}_{sanitized_category}_{max_pages}_{timestamp}.csv"
                            csv_path = os.path.join(SAVED_RUNS_DIR, csv_filename)
                            enriched_df.to_csv(csv_path, index=False)
                            st.info(f"📁 Enriched data exported to CSV at: {csv_path}")

                            # Save to session_state with run_label
                            run_label = f"Enriched_{sanitized_search_query}_{sanitized_category}_{max_pages}_{timestamp}"
                            st.session_state["saved_data"][run_label] = csv_path
                    except Exception as e:
                        st.error(f"❌ Error during enrichment: {e}")

    # New Section: Select Class Data Files
    st.markdown("---")
    st.subheader("📁 Select Class Data Files")

    class_data_files = get_class_data_files()
    if class_data_files:
        selected_class_files = st.multiselect("✅ Select files for classification:", options=class_data_files)
        if selected_class_files:
            if st.button("📥 Load Selected Class Data Files"):
                loaded_files = []
                for file in selected_class_files:
                    file_path = os.path.join("/Users/macbook/Library/Mobile Documents/com~apple~CloudDocs/Professioneel/Coding projects/marktplaats/data/class_data", file)
                    try:
                        df_class = pd.read_csv(file_path)
                        # Example: Merge with enriched_df or perform specific operations
                        if "enriched_df" in st.session_state:
                            st.session_state["enriched_df"] = pd.concat([st.session_state["enriched_df"], df_class], ignore_index=True)
                        else:
                            st.session_state["enriched_df"] = df_class
                        loaded_files.append(file)
                    except Exception as e:
                        st.error(f"❌ Error loading {file}: {e}")
                if loaded_files:
                    st.success(f"✅ Loaded class data files: {', '.join(loaded_files)}")
    else:
        st.warning("⚠️ No class data files available in the folder.")

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   TAB 3: ANALYSIS
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def tab_3_analysis():
    """
    Code originally found under 'with tabs[2]:'
    """
    st.header("📊 Analysis")

    # Retrieve the enriched DataFrame or preprocessed DataFrame if enriched is not available
    df_enriched = st.session_state.get("enriched_df", None)
    df_preprocessed = st.session_state.get("preprocessed_df", None)
    
    if df_enriched is None and df_preprocessed is None:
        st.warning("Please load and preprocess data first before analysis.")
    else:
        # Use enriched data if available, otherwise use preprocessed data
        df_analysis = df_enriched if df_enriched is not None else df_preprocessed
        is_enriched = df_enriched is not None

        # Sidebar Filters within the Analysis Tab
        st.sidebar.header("🔍 Filter Options")
        
        # Brand Selection (if enriched data available)
        if is_enriched and 'Brand' in df_analysis.columns:
            available_brands = sorted(df_analysis['Brand'].dropna().unique())
            selected_brands = st.sidebar.multiselect(
                "🏢 Select Brands:",
                options=available_brands,
                default=available_brands
            )
        else:
            selected_brands = None
        
        # Usage Selection (available in both preprocessed and enriched data)
        available_usages = sorted(df_analysis['Usage'].dropna().unique())
        selected_usages = st.sidebar.multiselect(
            "🔧 Select Usage Categories:",
            options=available_usages,
            default=available_usages
        )
        
        # Date Range Selection
        selected_date_range = st.sidebar.date_input(
            "📅 Select Date Range:",
        )
        
        # Price Range Selection
        price_col = 'CleanedPrice' if 'CleanedPrice' in df_analysis.columns else 'Price'
        
        # Handle non-numeric values in price column
        numeric_prices = pd.to_numeric(filtered_df[price_col], errors='coerce')
        valid_prices = numeric_prices.dropna()

        if valid_prices.empty:
            # No valid prices => Provide a fallback or skip the slider
            st.warning("No valid numeric prices available for filtering.")
            selected_price_range = (0.0, 0.0)  # or any fallback logic
        else:
            min_price = float(valid_prices.min())
            max_price = float(valid_prices.max())
            selected_price_range = st.sidebar.slider(
                "💰 Price Range",
                min_value=min_price,
                max_value=max_price,
                value=(min_price, max_price)
            )

        
        # Apply date filter if provided
        if isinstance(selected_date_range, tuple) and len(selected_date_range) == 2:
            start_date, end_date = selected_date_range
            filtered_df = filtered_df[
                (filtered_df['Date'] >= pd.to_datetime(start_date)) &
                (filtered_df['Date'] <= pd.to_datetime(end_date))
            ]
        
        st.markdown(f"### 📊 Selected Data: {len(filtered_df)} Listings")
        st.dataframe(filtered_df.head(10))

        # Analysis tabs - show all for enriched data, limited for preprocessed
        if is_enriched:
            analysis_tabs = st.tabs(["📈 Basic Analysis", "📝 Sentiment Analysis", "🏆 Best Deal", "🔍 Market Insights", "📊 Price Trends"])
        else:
            analysis_tabs = st.tabs(["📈 Basic Analysis"])

        ######################
        # Basic Analysis Tab 
        ######################
        with analysis_tabs[0]:
            st.subheader("📈 Basic Analysis")
            
            # Price Statistics
            st.write("### 💰 Price Analysis")
            numeric_prices = pd.to_numeric(filtered_df[price_col], errors='coerce')
            price_stats = numeric_prices.describe()
            st.write(price_stats)
            
            # Price Distribution Plot
            fig_price = px.histogram(filtered_df[pd.to_numeric(filtered_df[price_col], errors='coerce').notna()], 
                                    x=price_col, 
                                    title="Price Distribution")
            st.plotly_chart(fig_price)
            
            # Usage Distribution
            st.write("### 🔧 Usage Distribution")
            usage_counts = filtered_df['Usage'].value_counts()
            fig_usage = px.pie(values=usage_counts.values, names=usage_counts.index, title="Usage Distribution")
            st.plotly_chart(fig_usage)
            
            # Price Trends Over Time
            st.write("### 📅 Price Trends Over Time")
            trend_df = filtered_df[pd.to_numeric(filtered_df[price_col], errors='coerce').notna()]
            fig_trend = px.scatter(trend_df, x='Date', y=price_col, title="Price Trends Over Time")
            st.plotly_chart(fig_trend)

        ######################
        # Sentiment Analysis Tab (only if is_enriched)
        ######################
        if is_enriched:
            with analysis_tabs[1]:
                st.subheader("📝 Sentiment Analysis")
                if st.button("Analyze Sentiment"):
                    with st.spinner("🧠 Performing sentiment analysis..."):
                        try:
                            enriched_df = asyncio.run(sentiment_analysis_agent(filtered_df))
                            st.success("✅ Sentiment analysis complete.")
                            st.dataframe(enriched_df[['Title', 'Sentiment', 'SentimentScore']].head(10))
                        except Exception as e:
                            st.error(f"❌ Error during sentiment analysis: {e}")

            ######################
            # Price Trend Analysis Tab
            ######################
            with analysis_tabs[4]:
                st.subheader("📊 Price Trend Analysis")
                
                # Get price trend analysis
                try:
                    trend_analysis = asyncio.run(price_trend_analysis(filtered_df))
                    st.markdown(trend_analysis)
                    
                    # Create time series plot
                    if 'Date' in filtered_df.columns and 'CleanedPrice' in filtered_df.columns:
                        filtered_df['Date'] = pd.to_datetime(filtered_df['Date'])
                        filtered_df = filtered_df.sort_values('Date')
                        fig = px.line(filtered_df, x='Date', y='CleanedPrice', title='Price Trends Over Time')
                        st.plotly_chart(fig)
                except Exception as e:
                    st.error(f"❌ Error during price trend analysis: {e}")
                    
            ######################
            # Best Deal Tab
            ######################
            with analysis_tabs[2]:
                st.subheader("🏆 Best Deal Recommendation")
                if st.button("Get Best Deal"):
                    with st.spinner("🔍 Identifying the best deal..."):
                        try:
                            best_deal = asyncio.run(best_deal_recommendation(filtered_df))
                            st.info(f"**Best Deal Recommendation:**\n{best_deal}")
                        except Exception as e:
                            st.error(f"❌ Error during best deal recommendation: {e}")

            ######################
            # Market Insights Tab
            ######################
            with analysis_tabs[3]:
                st.subheader("🔍 Market Insights")
                if st.button("Generate Market Insights"):
                    with st.spinner("🧠 Generating market insights..."):
                        try:
                            insights = asyncio.run(market_insights_analysis(filtered_df))
                            st.session_state["insights"] = insights
                            st.info(f"**Market Insights:**\n{insights}")
                        except Exception as e:
                            st.error(f"❌ Error during market insights analysis: {e}")

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   TAB 4: EMAIL INSIGHTS
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def tab_4_email_insights():
    """
    Code originally found under 'with tabs[3]:'
    """
    st.subheader("📧 Email Insights")
    if st.button("Send Insights via Email"):
        insights = st.session_state.get("insights", "")
        if not insights:
            st.warning("⚠️ No insights available to send via email.")
        else:
            # Retrieve API key from secrets
            brevo_api_key = os.getenv("BREVO_API_KEY")
            # Define the API endpoint
            url = "https://api.brevo.com/v3/smtp/email"

            # Define the sender
            sender = {
                "name": "Marktplaats Scraper",
                "email": "davidkakaniss@gmail.com"
            }

            # Let user input recipient details
            st.subheader("Email Recipient Details")
            recipient_email = st.text_input("Recipient Email", value="david.kakanis@hotmail.com")
            recipient_name = st.text_input("Recipient Name", value="John Doe")

            if not recipient_email:
                st.warning("Please enter the recipient's email address.")
                return

            # Define the email content
            html_content = generate_email_html(insights)

            # Define the payload
            payload = {
                "sender": sender,
                "to": [
                    {
                        "email": recipient_email,
                        "name": recipient_name
                    }
                ],
                "subject": "Market Insights from Marktplaats Scraper",
                "htmlContent": html_content
            }

            # Define the headers
            headers = {
                "accept": "application/json",
                "api-key": brevo_api_key,
                "content-type": "application/json"
            }

            # Send the POST request
            try:
                response = requests.post(url, headers=headers, json=payload)
                response.raise_for_status()  # Raise an exception for HTTP errors
                st.success("Email sent successfully!")
            except requests.exceptions.HTTPError as http_err:
                st.error(f"HTTP error occurred: {http_err} - {response.text}")
            except Exception as err:
                st.error(f"An error occurred: {err}")

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   TAB 5: COMPREHENSIVE ANALYSIS
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def tab_5_comprehensive_analysis():
    """
    Code originally found under 'with tabs[4]:'
    """
    st.header("🔍 Comprehensive Analysis")

    # Sub-tabs for different comprehensive analyses
    comprehensive_tabs = st.tabs(["📈 Market Demand Forecasting", "🔄 Detailed Comparative Analysis", "🤖 Agent-Based Analysis"])

    # Sub-tab 1: Market Demand Forecasting
    with comprehensive_tabs[0]:
        st.subheader("📈 Market Demand Forecasting")
        if st.button("Forecast Market Demand"):
            df_for_forecasting = st.session_state.get("enriched_df", pd.DataFrame())
            if df_for_forecasting.empty:
                st.warning("⚠️ No data to analyze. Please enrich data first.")
            else:
                with st.spinner("📈 Forecasting market demand..."):
                    try:
                        forecasting = asyncio.run(market_demand_forecasting(df_for_forecasting))
                        st.info(f"**Market Demand Forecasting:**\n{forecasting}")
                    except Exception as e:
                        st.error(f"❌ Error during market demand forecasting: {e}")

    # Sub-tab 2: Detailed Comparative Analysis
    with comprehensive_tabs[1]:
        st.subheader("🔄 Detailed Comparative Analysis")
        if st.button("Compare Listings"):
            df_for_comparison = st.session_state.get("enriched_df", pd.DataFrame())
            if df_for_comparison.empty:
                st.warning("⚠️ No data to analyze. Please enrich data first.")
            else:
                with st.spinner("🔍 Performing detailed comparative analysis..."):
                    try:
                        comparison = asyncio.run(detailed_comparative_analysis(df_for_comparison))
                        st.info(f"**Comparative Analysis:**\n\n{comparison}")
                    except Exception as e:
                        st.error(f"❌ Error during comparative analysis: {e}")

    # Sub-tab 3: Agent-Based Analysis
    with comprehensive_tabs[2]:
        st.subheader("🤖 Agent-Based Analysis")
        if st.button("Run Comprehensive Agent Analysis"):
            df_for_agent = st.session_state.get("enriched_df", pd.DataFrame())
            if df_for_agent.empty:
                st.warning("⚠️ No data to analyze. Please enrich data first.")
            else:
                with st.spinner("🤖 Running comprehensive agent analysis..."):
                    try:
                        agent_insights = asyncio.run(run_agent_analysis(df_for_agent))
                        st.info(f"**Comprehensive Agent Analysis:**\n\n{agent_insights}")
                    except Exception as e:
                        st.error(f"❌ Error during comprehensive agent analysis: {e}")

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   TAB 6: DASHBOARD
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def tab_6_dashboard():
    """
    Code originally found under 'with tabs[5]:'
    """
    st.header("📊 Dashboard")

    df_dash = st.session_state.get("enriched_df", pd.DataFrame())
    if df_dash.empty:
        st.warning("⚠️ No data to analyze. Please enrich data first.")
    else:
        show_dashboard(df_dash)

# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   TAB 7: BACKGROUND ANALYSIS
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def tab_7_background_analysis():
    """
    Code originally found under 'with tabs[6]:'
    """
    st.header("Background Analysis and Email Alerts")
    queue_background_analysis()
    show_analysis_queue_status()

def tab_8_image_search():
    """
    Code for the Image Search tab
    """
    st.header("Image Search")
    image_search.find_similar_products()  # Call the specific function instead of the module



# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
#   DOWNLOAD REPORT FUNCTION (UNCHANGED)
# ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
def download_analysis_report(df: pd.DataFrame):
    if df.empty:
        st.warning("No data available to generate a report.")
        return

    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading('Marktplaats Scraper Analysis Report', 0)

    # Add data summary
    doc.add_heading('Data Summary', 1)
    doc.add_paragraph(f"Number of listings: {len(df)}")

    # Add visualizations
    if 'Brand' in df.columns and 'CleanedPrice' in df.columns:
        doc.add_heading('Brand Prices', 1)
        brand_prices = df.groupby('Brand')['CleanedPrice'].mean().reset_index()
        brand_prices = brand_prices.sort_values('CleanedPrice', ascending=False)
        data = [['Brand', 'Average Price']] + brand_prices.values.tolist()
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(data):
            if i == 0:
                # Header row
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)
            else:
                table.add_row()
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)

        fig, ax = plt.subplots(figsize=(10, 6))
        sns.barplot(x='Brand', y='CleanedPrice', data=brand_prices, ax=ax)
        ax.set_title('Average Brand Prices')
        ax.set_xlabel('Brand')
        ax.set_ylabel('Average Price (€)')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'Sentiment' in df.columns:
        doc.add_heading('Sentiment Distribution', 1)
        sentiment_counts = df['Sentiment'].value_counts().reset_index()
        sentiment_counts.columns = ['Sentiment', 'Count']
        data = sentiment_counts.values.tolist()
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(data):
            if i == 0:
                # Header row
                table.cell(i, 0).text = "Sentiment"
                table.cell(i, 1).text = "Count"
            else:
                table.add_row()
                table.cell(i, 0).text = str(row[0])
                table.cell(i, 1).text = str(row[1])

        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(sentiment_counts['Count'], labels=sentiment_counts['Sentiment'], autopct='%1.1f%%')
        ax.axis('equal')
        ax.set_title('Sentiment Distribution')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'CleanedPrice' in df.columns:
        doc.add_heading('Price Distribution', 1)
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.histplot(data=df, x='CleanedPrice', ax=ax)
        ax.set_title('Price Distribution')
        ax.set_xlabel('Price (€)')
        ax.set_ylabel('Number of Listings')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'Usage' in df.columns and 'CleanedPrice' in df.columns:
        doc.add_heading('Usage vs. Average Price', 1)
        usage_price = df.groupby('Usage')['CleanedPrice'].mean().reset_index()
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.scatterplot(x='Usage', y='CleanedPrice', data=usage_price, ax=ax)
        ax.set_title('Usage vs. Average Price')
        ax.set_xlabel('Usage')
        ax.set_ylabel('Average Price (€)')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    # Save the document
    file_name = f"analysis_report.docx"
    doc.save(file_name)

    # Download the document
    with open(file_name, "rb") as file:
        btn = st.download_button(
            label="Download Analysis Report",
            data=file.read(),
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )



if __name__ == "__main__":
    load_orchestrate_and_display()





