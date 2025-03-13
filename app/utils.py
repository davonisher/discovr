import re
import os
import pandas as pd
from datetime import datetime
from price_parser import Price
from typing import List, Optional
from playwright.async_api import Route
import io
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import streamlit as st

def clean_price(price_str: str) -> float:
    """
    Extracts the numeric amount from a price string using the price-parser library.

    Parameters:
    - price_str (str): Raw price string (e.g., "22,90 €", "$119.00").

    Returns:
    - float: Cleaned price amount. Returns 0.0 if the price can't be parsed.
    """
    if not isinstance(price_str, str):
        return 0.0

    # Use price-parser to parse the price
    price = Price.fromstring(price_str)

    # Return the amount as a float, defaulting to 0.0 if parsing fails
    return price.amount_float or 0.0

def parse_date(date_str: str) -> Optional[datetime.date]:
    """
    Parse a date string into a datetime.date object.
    Returns None if parsing fails.
    """
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return None

def sanitize_filename_component(component: str) -> str:
    """
    Removes or replaces characters that are invalid in filenames.
    """
    return "".join(c for c in component if c.isalnum() or c in (' ', '_', '-')).rstrip()

def is_valid_email(email: str) -> bool:
    """
    Validates the email format.
    """
    regex = r'^[\w\.-]+@[\w\.-]+\.\w+$'
    return re.match(regex, email) is not None


def get_category_options() -> List[str]:
    """
    Get a list of available category options from the categories file.
    """
    try:
        categories_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), 
            "modules", "scrapers", "cat.txt"
        )
        with open(categories_path, "r") as f:
            return [line.strip() for line in f if line.strip()]
    except FileNotFoundError:
        # Fallback to default categories if file not found
        return [
            "audio-tv-en-foto",
            "computers-en-software",
            "consoles-en-games",
            "telefoons",
            "sport-en-fitness"
        ]

async def block_aggressively(route: Route):
    """
    Prevent the browser from downloading images, CSS, etc.
    This speeds up scraping if you only need the HTML.
    """
    if route.request.resource_type != "document":
        await route.abort()
    else:
        await route.continue_()

def load_user_agents() -> List[str]:
    """
    Load user agents from a text file, one per line.
    """
    user_agents_path = os.path.join(os.path.dirname(os.path.dirname(__file__)),    "modules", "scrapers", "user-agents.txt")
    try:
        with open(user_agents_path, "r") as f:
            return [line.strip() for line in f if line.strip()]
    except FileNotFoundError:
        # Return a default user agent if file not found
        return ["Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"]

def group_prices_by_usage(df: pd.DataFrame) -> pd.DataFrame:
    """
    Groups prices by usage and calculates average price per usage category.
    
    Parameters:
    - df (pd.DataFrame): DataFrame containing 'Usage' and 'CleanedPrice' columns
    
    Returns:
    - pd.DataFrame: Grouped DataFrame with average prices per usage category
    """
    if 'Usage' not in df.columns or 'CleanedPrice' not in df.columns:
        return pd.DataFrame()
    grouped = df.groupby('Usage', dropna=False)['CleanedPrice'].mean().reset_index()
    grouped.rename(columns={'CleanedPrice': 'AvgPrice'}, inplace=True)
    return grouped 

def download_analysis_report(df: pd.DataFrame):
    """
    Generates and allows downloading of a comprehensive analysis report in Word format.
    """
    if df.empty:
        st.warning("No data available to generate a report.")
        return

    # Create a new Word document
    doc = Document()

    # Add a title
    doc.add_heading('Marktplaats Scraper Analysis Report', 0)

    # Add data summary
    doc.add_heading('Data Summary', 1)
    doc.add_paragraph(f"Number of listings: {len(df)}")

    # Add visualizations
    if 'Brand' in df.columns and 'CleanedPrice' in df.columns:
        doc.add_heading('Brand Prices', 1)
        brand_prices = df.groupby('Brand')['CleanedPrice'].mean().reset_index()
        brand_prices = brand_prices.sort_values('CleanedPrice', ascending=False)
        data = [['Brand', 'Average Price']] + brand_prices.values.tolist()
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(data):
            if i == 0:
                # Header row
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)
            else:
                table.add_row()
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)

        fig, ax = plt.subplots(figsize=(10, 6))
        sns.barplot(x='Brand', y='CleanedPrice', data=brand_prices, ax=ax)
        ax.set_title('Average Brand Prices')
        ax.set_xlabel('Brand')
        ax.set_ylabel('Average Price (€)')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'Sentiment' in df.columns:
        doc.add_heading('Sentiment Distribution', 1)
        sentiment_counts = df['Sentiment'].value_counts().reset_index()
        sentiment_counts.columns = ['Sentiment', 'Count']
        data = sentiment_counts.values.tolist()
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(data):
            if i == 0:
                # Header row
                table.cell(i, 0).text = "Sentiment"
                table.cell(i, 1).text = "Count"
            else:
                table.add_row()
                table.cell(i, 0).text = str(row[0])
                table.cell(i, 1).text = str(row[1])

        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(sentiment_counts['Count'], labels=sentiment_counts['Sentiment'], autopct='%1.1f%%')
        ax.axis('equal')
        ax.set_title('Sentiment Distribution')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'CleanedPrice' in df.columns:
        doc.add_heading('Price Distribution', 1)
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.histplot(data=df, x='CleanedPrice', ax=ax)
        ax.set_title('Price Distribution')
        ax.set_xlabel('Price (€)')
        ax.set_ylabel('Number of Listings')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    if 'Usage' in df.columns and 'CleanedPrice' in df.columns:
        doc.add_heading('Usage vs. Average Price', 1)
        usage_price = df.groupby('Usage')['CleanedPrice'].mean().reset_index()
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.scatterplot(x='Usage', y='CleanedPrice', data=usage_price, ax=ax)
        ax.set_title('Usage vs. Average Price')
        ax.set_xlabel('Usage')
        ax.set_ylabel('Average Price (€)')
        img_data = io.BytesIO()
        fig.savefig(img_data, format='png', dpi=300)
        img_data.seek(0)
        doc.add_picture(img_data, width=Inches(6))

    # Save the document
    file_name = f"analysis_report.docx"
    doc.save(file_name)

    # Download the document
    with open(file_name, "rb") as file:
        btn = st.download_button(
            label="Download Analysis Report",
            data=file.read(),
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def clean_price(price_str: str) -> float:
    """
    Extracts the numeric amount from a price string using the price-parser library.

    Parameters:
    - price_str (str): Raw price string (e.g., "22,90 €", "$119.00").

    Returns:
    - float: Cleaned price amount. Returns 0.0 if the price can't be parsed.
    """
    if not isinstance(price_str, str):
        return 0.0

    # Use price-parser to parse the price
    price = Price.fromstring(price_str)

    # Return the amount as a float, defaulting to 0.0 if parsing fails
    return price.amount_float or 0.0

def parse_date(date_str: str) -> Optional[datetime.date]:
    """
    Parse a date string into a datetime.date object.
    Returns None if parsing fails.
    """
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception:
        return None



def get_class_data_files() -> List[str]:
    """
    Get a list of available class data files.
    """
    class_data_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "class_data")
    if not os.path.exists(class_data_dir):
        return []
    return [f for f in os.listdir(class_data_dir) if f.endswith('.csv')]

def get_raw_data_files() -> List[str]:
    """
    Get a list of available raw data files.
    """
    raw_data_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "raw_data")
    if not os.path.exists(raw_data_dir):
        return []
    return [f for f in os.listdir(raw_data_dir) if f.endswith('.csv')]

def get_category_options() -> List[str]:
    """
    Get a list of available category options from the categories file.
    """
    try:
        categories_path = os.path.join(
            os.path.dirname(os.path.dirname(__file__)), 
            "modules", "scrapers", "cat.txt"
        )
        with open(categories_path, "r") as f:
            return [line.strip() for line in f if line.strip()]
    except FileNotFoundError:
        # Fallback to default categories if file not found
        return [
            "audio-tv-en-foto",
            "computers-en-software",
            "consoles-en-games",
            "telefoons",
            "sport-en-fitness"
        ]

def preprocess_enriched_df(df: pd.DataFrame) -> pd.DataFrame:
    """
    Preprocesses the enriched DataFrame by converting date columns to datetime
    and cleaning price values.
    
    Parameters:
    - df (pd.DataFrame): The enriched DataFrame.
    
    Returns:
    - pd.DataFrame: The preprocessed DataFrame.
    """
    if 'Date' in df.columns:
        df['ClassifiedDate'] = df['Date'].apply(parse_date)
    
    if 'Price' in df.columns:
        df['CleanedPrice'] = df['Price'].apply(clean_price)
    
    return df 



def group_prices_by_usage(df: pd.DataFrame) -> pd.DataFrame:
    """
    Groups prices by usage and calculates average price per usage category.
    
    Parameters:
    - df (pd.DataFrame): DataFrame containing 'Usage' and 'CleanedPrice' columns
    
    Returns:
    - pd.DataFrame: Grouped DataFrame with average prices per usage category
    """
    if 'Usage' not in df.columns or 'CleanedPrice' not in df.columns:
        return pd.DataFrame()
    grouped = df.groupby('Usage', dropna=False)['CleanedPrice'].mean().reset_index()
    grouped.rename(columns={'CleanedPrice': 'AvgPrice'}, inplace=True)
    return grouped 
